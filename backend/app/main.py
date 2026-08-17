import json
import ipaddress
import hashlib
import hmac
import os
import re
import socket
import ssl
import subprocess
import sys
import time
import uuid
import zipfile
from io import BytesIO
from concurrent.futures import ThreadPoolExecutor
from datetime import datetime, timezone
from html import escape
from html.parser import HTMLParser
from pathlib import Path
from typing import Literal
from urllib.error import HTTPError, URLError
from urllib.parse import quote, urlparse
from urllib.request import Request, urlopen

from cryptography.fernet import Fernet, InvalidToken
import certifi
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from fastapi import BackgroundTasks, Depends, FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, Response
import httpx
from pydantic import BaseModel, Field, field_validator
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from markdown_it import MarkdownIt
from markdown_it.tree import SyntaxTreeNode
from sqlalchemy import DateTime, Integer, String, Text, create_engine, func, select, text
from sqlalchemy.orm import DeclarativeBase, Mapped, Session, mapped_column, sessionmaker

DB_PATH = Path(__file__).resolve().parent.parent / "geo.db"
engine = create_engine(f"sqlite:///{DB_PATH}", connect_args={"check_same_thread": False})
SessionLocal = sessionmaker(bind=engine, autoflush=False)
DEFAULT_CONSTRAINT = """1. 开头用金字塔原理写总体摘要，100-300 字，要有精华和干货
2. 重要结论有来源或可验证证据，体现品牌/主题，或有独立结论
3. 提供新的事实、数据、案例、方法或判断标准
4. 主打 2 个关键词，不要同时抢十几个互不相关的关键词
5. 适当介绍目标人群
6. 不要大量使用模糊代词
7. 不要标题党，要干货，要真诚"""
FAQ_CATEGORIES = ["采购流程", "产品描述与认证", "价格与起订量", "售后与质保"]


class Base(DeclarativeBase):
    pass


class Brand(Base):
    __tablename__ = "brands"
    id: Mapped[int] = mapped_column(primary_key=True)
    name: Mapped[str] = mapped_column(String(100), unique=True)
    company: Mapped[str] = mapped_column(String(160))
    alias: Mapped[str] = mapped_column(String(120), default="")
    region: Mapped[str] = mapped_column(String(100), default="")
    constraint: Mapped[str] = mapped_column(Text, default="")
    brand_code: Mapped[str] = mapped_column(String(20), default="")
    founded: Mapped[str] = mapped_column(String(20), default="")
    products: Mapped[str] = mapped_column(String(240), default="")
    target_customers: Mapped[str] = mapped_column(String(240), default="")
    core_capability: Mapped[str] = mapped_column(String(240), default="")
    certifications: Mapped[str] = mapped_column(Text, default="")
    created_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.now)


class Faq(Base):
    __tablename__ = "faqs"
    id: Mapped[int] = mapped_column(primary_key=True)
    faq_code: Mapped[str] = mapped_column(String(20), default="")
    category: Mapped[str] = mapped_column(String(80))
    question: Mapped[str] = mapped_column(String(300))
    answer: Mapped[str] = mapped_column(Text)
    created_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.now)


class Knowledge(Base):
    __tablename__ = "knowledge"
    id: Mapped[int] = mapped_column(primary_key=True)
    name: Mapped[str] = mapped_column(String(160))
    category: Mapped[str] = mapped_column(String(80))
    content: Mapped[str] = mapped_column(Text)
    created_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.now)


class AuditLog(Base):
    __tablename__ = "audit_logs"
    id: Mapped[int] = mapped_column(primary_key=True)
    category: Mapped[str] = mapped_column(String(20), default="其他")
    action: Mapped[str] = mapped_column(String(80))
    detail: Mapped[str] = mapped_column(String(300), default="")
    created_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.now)


class GeneratedArticle(Base):
    __tablename__ = "generated_articles"
    id: Mapped[int] = mapped_column(primary_key=True)
    article_code: Mapped[str] = mapped_column(String(20), unique=True, index=True)
    input_text: Mapped[str] = mapped_column(Text)
    supplemental_prompt: Mapped[str] = mapped_column(Text, default="")
    content: Mapped[str] = mapped_column(Text)
    brand_id: Mapped[int] = mapped_column(Integer, default=0)
    brand_name: Mapped[str] = mapped_column(String(100), default="")
    include_faq: Mapped[int] = mapped_column(Integer, default=1)
    include_knowledge: Mapped[int] = mapped_column(Integer, default=1)
    image_prompt: Mapped[str] = mapped_column(Text, default="")
    selected_images: Mapped[str] = mapped_column(Text, default="[]")
    created_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.now)
    content_saved_at: Mapped[datetime | None] = mapped_column(DateTime, nullable=True)
    confirmed_at: Mapped[datetime | None] = mapped_column(DateTime, nullable=True)


class TranslationTask(Base):
    __tablename__ = "translation_tasks"
    id: Mapped[int] = mapped_column(primary_key=True)
    article_id: Mapped[int] = mapped_column(Integer, index=True)
    language: Mapped[str] = mapped_column(String(80))
    status: Mapped[str] = mapped_column(String(20), default="processing")
    progress: Mapped[int] = mapped_column(Integer, default=8)
    error: Mapped[str] = mapped_column(Text, default="")
    created_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.now)
    completed_at: Mapped[datetime | None] = mapped_column(DateTime, nullable=True)


class GeoReport(Base):
    __tablename__ = "geo_reports"
    id: Mapped[int] = mapped_column(primary_key=True)
    report_code: Mapped[str] = mapped_column(String(30), unique=True, index=True)
    keyword: Mapped[str] = mapped_column(String(300), index=True)
    provider_names: Mapped[str] = mapped_column(Text, default="[]")
    result_json: Mapped[str] = mapped_column(Text)
    created_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.now)


class ModelProvider(Base):
    __tablename__ = "model_providers"
    id: Mapped[int] = mapped_column(primary_key=True)
    name: Mapped[str] = mapped_column(String(100), unique=True)
    base_url: Mapped[str] = mapped_column(String(300))
    model: Mapped[str] = mapped_column(String(160))
    provider_type: Mapped[str] = mapped_column(String(30), default="qwen")
    protocol: Mapped[str] = mapped_column(String(30), default="responses")
    api_key_encrypted: Mapped[str] = mapped_column(Text, default="")
    api_secret_encrypted: Mapped[str] = mapped_column(Text, default="")
    aux_api_key_encrypted: Mapped[str] = mapped_column(Text, default="")
    enabled: Mapped[int] = mapped_column(Integer, default=1)
    created_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.now)


class MediaProvider(Base):
    __tablename__ = "media_providers"
    id: Mapped[int] = mapped_column(primary_key=True)
    name: Mapped[str] = mapped_column(String(100), unique=True)
    media_type: Mapped[str] = mapped_column(String(20), default="image")
    base_url: Mapped[str] = mapped_column(String(300), default="https://ark.cn-beijing.volces.com/api/v3")
    model: Mapped[str] = mapped_column(String(160))
    api_key_encrypted: Mapped[str] = mapped_column(Text, default="")
    enabled: Mapped[int] = mapped_column(Integer, default=1)
    created_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.now)


class SearchProvider(Base):
    __tablename__ = "search_providers"
    id: Mapped[int] = mapped_column(primary_key=True)
    name: Mapped[str] = mapped_column(String(100), unique=True)
    base_url: Mapped[str] = mapped_column(String(300))
    api_key_encrypted: Mapped[str] = mapped_column(Text, default="")
    search_depth: Mapped[str] = mapped_column(String(20), default="advanced")
    max_results: Mapped[int] = mapped_column(Integer, default=10)
    enabled: Mapped[int] = mapped_column(Integer, default=1)
    created_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.now)


class OfficialMonitorTask(Base):
    __tablename__ = "official_monitor_tasks"
    id: Mapped[int] = mapped_column(primary_key=True)
    task_code: Mapped[str] = mapped_column(String(30), unique=True, index=True)
    keyword: Mapped[str] = mapped_column(String(300))
    platform: Mapped[str] = mapped_column(String(40), index=True)
    status: Mapped[str] = mapped_column(String(20), default="pending")
    error_message: Mapped[str] = mapped_column(Text, default="")
    note: Mapped[str] = mapped_column(String(240), default="")
    started_at: Mapped[datetime | None] = mapped_column(DateTime, nullable=True)
    finished_at: Mapped[datetime | None] = mapped_column(DateTime, nullable=True)
    created_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.now)


class OfficialMonitorAnswer(Base):
    __tablename__ = "official_monitor_answers"
    id: Mapped[int] = mapped_column(primary_key=True)
    task_id: Mapped[int] = mapped_column(Integer, index=True)
    answer_text: Mapped[str] = mapped_column(Text, default="")
    answer_html: Mapped[str] = mapped_column(Text, default="")
    raw_capture_json: Mapped[str] = mapped_column(Text, default="{}")
    screenshot_path: Mapped[str] = mapped_column(String(400), default="")
    created_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.now)


class OfficialMonitorSource(Base):
    __tablename__ = "official_monitor_sources"
    id: Mapped[int] = mapped_column(primary_key=True)
    task_id: Mapped[int] = mapped_column(Integer, index=True)
    title: Mapped[str] = mapped_column(String(400), default="")
    url: Mapped[str] = mapped_column(Text, default="")
    normalized_url: Mapped[str] = mapped_column(Text, default="")
    domain: Mapped[str] = mapped_column(String(200), default="")
    root_domain: Mapped[str] = mapped_column(String(200), default="")
    source_type: Mapped[str] = mapped_column(String(40), default="unknown")
    rank_index: Mapped[int] = mapped_column(Integer, default=0)
    display_text: Mapped[str] = mapped_column(Text, default="")
    created_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.now)


class OfficialMonitorRun(Base):
    __tablename__ = "official_monitor_runs"
    id: Mapped[int] = mapped_column(primary_key=True)
    task_id: Mapped[int] = mapped_column(Integer, index=True)
    step: Mapped[str] = mapped_column(String(60))
    status: Mapped[str] = mapped_column(String(20), default="info")
    message: Mapped[str] = mapped_column(Text, default="")
    created_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.now)


class BrandInput(BaseModel):
    name: str
    company: str
    alias: str = ""
    region: str = ""
    founded: str = ""
    products: str = ""
    target_customers: str = ""
    core_capability: str = ""
    certifications: str = ""
    constraint: str = DEFAULT_CONSTRAINT


class FaqInput(BaseModel):
    category: str
    question: str
    answer: str

    @field_validator("question", "answer")
    @classmethod
    def validate_required_text(cls, value: str):
        if not value.strip():
            raise ValueError("提问和答案不能为空")
        return value.strip()

    @field_validator("category")
    @classmethod
    def validate_category(cls, value: str):
        if value not in FAQ_CATEGORIES:
            raise ValueError("请选择有效的 FAQ 类型")
        return value


class KnowledgeInput(BaseModel):
    name: str
    category: str
    content: str


class ComposeInput(BaseModel):
    prompt: str
    supplemental_prompt: str = ""
    brand_id: int | None = None
    include_faq: bool = True
    include_knowledge: bool = True

    @field_validator("prompt")
    @classmethod
    def validate_prompt(cls, value: str):
        if not value.strip():
            raise ValueError("输入文案不能为空")
        return value.strip()


class ArticleDraftInput(BaseModel):
    content: str
    image_prompt: str = ""
    selected_images: list[str] = Field(default_factory=list)


class ArticleOptimizeInput(BaseModel):
    instruction: str

    @field_validator("instruction")
    @classmethod
    def optimize_instruction_is_not_empty(cls, value: str):
        if not value.strip():
            raise ValueError("请输入优化提示词")
        return value.strip()


class TranslationInput(BaseModel):
    language: Literal["English", "Bahasa Indonesia", "日语"]


class GeoInput(BaseModel):
    keyword: str
    provider_ids: list[int]

    @field_validator("keyword")
    @classmethod
    def keyword_is_not_empty(cls, value: str):
        if not value.strip():
            raise ValueError("品牌词或关键词不能为空")
        return value.strip()

    @field_validator("provider_ids")
    @classmethod
    def providers_are_valid(cls, value: list[int]):
        if not value:
            raise ValueError("请至少选择一个模型通道")
        return list(dict.fromkeys(value))


class GeoReportEditInput(BaseModel):
    keyword: str
    summaries: dict[str, str] = Field(default_factory=dict)

    @field_validator("keyword")
    @classmethod
    def report_keyword_is_not_empty(cls, value: str):
        if not value.strip():
            raise ValueError("报告关键词不能为空")
        return value.strip()

class ModelProviderInput(BaseModel):
    name: str
    base_url: str
    model: str
    provider_type: Literal["qwen", "hunyuan", "volcengine", "tencent_search", "deepseek"] = "qwen"
    protocol: Literal["responses", "chat_completions"] = "chat_completions"
    api_key: str = ""
    api_secret: str = ""
    aux_api_key: str = ""
    enabled: bool = True

    @field_validator("name", "base_url", "model")
    @classmethod
    def provider_fields_are_not_empty(cls, value: str):
        if not value.strip():
            raise ValueError("该字段不能为空")
        return value.strip()


class MediaProviderInput(BaseModel):
    name: str
    media_type: Literal["image", "video"]
    base_url: str = "https://ark.cn-beijing.volces.com/api/v3"
    model: str
    api_key: str = ""
    enabled: bool = True

    @field_validator("name", "base_url", "model")
    @classmethod
    def media_fields_are_not_empty(cls, value: str):
        if not value.strip():
            raise ValueError("该字段不能为空")
        return value.strip()


class SearchProviderInput(BaseModel):
    name: str
    base_url: str
    api_key: str = ""
    search_depth: Literal["basic", "advanced"] = "advanced"
    max_results: int = 10
    enabled: bool = True

    @field_validator("name", "base_url")
    @classmethod
    def search_provider_fields_are_not_empty(cls, value: str):
        if not value.strip():
            raise ValueError("该字段不能为空")
        return value.strip()

    @field_validator("max_results")
    @classmethod
    def max_results_in_range(cls, value: int):
        if value < 1 or value > 20:
            raise ValueError("搜索结果数量需在 1 到 20 之间")
        return value


class ConstraintInput(BaseModel):
    constraint: str


class LogEventInput(BaseModel):
    category: Literal["新增", "修改", "生成", "查询", "删除"]
    action: str
    detail: str = ""


class OfficialMonitorTaskInput(BaseModel):
    keyword: str
    platforms: list[Literal["chatgpt", "deepseek", "doubao", "hunyuan", "claude"]]
    note: str = ""

    @field_validator("keyword")
    @classmethod
    def official_keyword_not_empty(cls, value: str):
        if not value.strip():
            raise ValueError("监测关键词不能为空")
        return value.strip()

    @field_validator("platforms")
    @classmethod
    def official_platforms_not_empty(cls, value: list[str]):
        if not value:
            raise ValueError("请至少选择一个平台")
        return list(dict.fromkeys(value))


app = FastAPI(title="GEO Operations API", version="0.1.0")
app.add_middleware(CORSMiddleware, allow_origins=["http://localhost:5173", "http://127.0.0.1:5173"], allow_methods=["*"], allow_headers=["*"])


def db_session():
    with SessionLocal() as session:
        yield session


def serialize(model):
    result = {column.name: getattr(model, column.name) for column in model.__table__.columns}
    for key, value in result.items():
        if isinstance(value, datetime):
            result[key] = value.isoformat()
    return result


def serialize_provider(provider: ModelProvider):
    result = serialize(provider)
    result.pop("api_key_encrypted", None)
    result.pop("api_secret_encrypted", None)
    result.pop("aux_api_key_encrypted", None)
    result["enabled"] = bool(provider.enabled)
    result["has_api_key"] = bool(provider.api_key_encrypted)
    result["has_api_secret"] = bool(provider.api_secret_encrypted)
    result["has_aux_api_key"] = bool(provider.aux_api_key_encrypted)
    return result


def serialize_media_provider(provider: MediaProvider):
    result = serialize(provider)
    result.pop("api_key_encrypted", None)
    result["enabled"] = bool(provider.enabled)
    result["has_api_key"] = bool(provider.api_key_encrypted)
    return result


def serialize_search_provider(provider: SearchProvider):
    result = serialize(provider)
    result.pop("api_key_encrypted", None)
    result["enabled"] = bool(provider.enabled)
    result["has_api_key"] = bool(provider.api_key_encrypted)
    return result


def serialize_official_task(task: OfficialMonitorTask):
    result = serialize(task)
    answer = session_get_answer(task.id)
    sources = session_get_sources(task.id)
    result["source_count"] = len(sources)
    result["page_count"] = len(sources)
    result["has_answer"] = bool(answer)
    return result


def session_get_answer(task_id: int, session: Session | None = None):
    owns_session = session is None
    if owns_session:
        session = SessionLocal()
    try:
        return session.scalar(select(OfficialMonitorAnswer).where(OfficialMonitorAnswer.task_id == task_id).order_by(OfficialMonitorAnswer.id.desc()))
    finally:
        if owns_session:
            session.close()


def session_get_sources(task_id: int, session: Session | None = None):
    owns_session = session is None
    if owns_session:
        session = SessionLocal()
    try:
        return session.scalars(select(OfficialMonitorSource).where(OfficialMonitorSource.task_id == task_id).order_by(OfficialMonitorSource.rank_index.asc(), OfficialMonitorSource.id.asc())).all()
    finally:
        if owns_session:
            session.close()


OFFICIAL_PLATFORM_META = {
    "chatgpt": {"label": "ChatGPT", "login_url": "https://chatgpt.com", "status": "manual_login_required"},
    "deepseek": {"label": "DeepSeek", "login_url": "https://chat.deepseek.com", "status": "manual_login_required"},
    "doubao": {"label": "豆包", "login_url": "https://www.doubao.com", "status": "planned"},
    "hunyuan": {"label": "混元", "login_url": "https://hunyuan.tencent.com", "status": "planned"},
    "claude": {"label": "Claude", "login_url": "https://claude.ai", "status": "planned"},
}
OFFICIAL_MONITOR_ARTIFACT_DIR = Path(__file__).resolve().parent.parent / "official_monitor_artifacts"
OFFICIAL_MONITOR_PROFILE_DIR = Path(__file__).resolve().parent.parent / "official_monitor_profiles"


def normalize_domain(url: str) -> tuple[str, str]:
    host = url.split("://", 1)[-1].split("/", 1)[0].lower()
    parts = [part for part in host.split(".") if part]
    root = ".".join(parts[-2:]) if len(parts) >= 2 else host
    return host, root


def source_type_for_domain(host: str) -> str:
    if host.endswith("mp.weixin.qq.com"):
        return "wechat"
    if host.endswith("baike.baidu.com") or "wikipedia.org" in host:
        return "encyclopedia"
    if "douyin.com" in host:
        return "short_video"
    if host.endswith(".gov.cn") or host.endswith(".gov"):
        return "government"
    if any(media in host for media in ["sina.cn", "ifeng.com", "sohu.com", "163.com", "qq.com"]):
        return "news_media"
    return "official_site" if host.count(".") >= 1 else "unknown"


def official_monitor_platforms():
    return [
        {"key": key, **value}
        for key, value in OFFICIAL_PLATFORM_META.items()
    ]


def append_monitor_run(session: Session, task_id: int, step: str, status: str, message: str):
    session.add(OfficialMonitorRun(task_id=task_id, step=step, status=status, message=message))


def run_official_monitor_worker(platform: str, keyword: str, task_code: str) -> dict:
    worker = Path(__file__).resolve().parent / "official_monitor_worker.py"
    OFFICIAL_MONITOR_ARTIFACT_DIR.mkdir(parents=True, exist_ok=True)
    OFFICIAL_MONITOR_PROFILE_DIR.mkdir(parents=True, exist_ok=True)
    command = [
        sys.executable,
        str(worker),
        "--platform", platform,
        "--keyword", keyword,
        "--task-code", task_code,
        "--profile-dir", str(OFFICIAL_MONITOR_PROFILE_DIR / platform),
        "--artifact-dir", str(OFFICIAL_MONITOR_ARTIFACT_DIR),
    ]
    completed = subprocess.run(command, capture_output=True, text=True, timeout=240)
    if completed.returncode != 0:
        raise RuntimeError((completed.stderr or completed.stdout or "官方监测 Worker 执行失败").strip())
    try:
        return json.loads(completed.stdout.strip() or "{}")
    except json.JSONDecodeError as error:
        raise RuntimeError("官方监测 Worker 返回格式无法解析") from error


def execute_official_monitor_task(task_id: int):
    with SessionLocal() as session:
        task = session.get(OfficialMonitorTask, task_id)
        if not task:
            return
        task.status = "running"
        task.error_message = ""
        task.started_at = datetime.now()
        session.query(OfficialMonitorAnswer).filter(OfficialMonitorAnswer.task_id == task.id).delete()
        session.query(OfficialMonitorSource).filter(OfficialMonitorSource.task_id == task.id).delete()
        session.query(OfficialMonitorRun).filter(OfficialMonitorRun.task_id == task.id).delete()
        append_monitor_run(session, task.id, "task_created", "info", "任务已创建，进入第一版监测流程。")
        platform_meta = OFFICIAL_PLATFORM_META[task.platform]
        append_monitor_run(session, task.id, "adapter", "info", f"平台适配器：{platform_meta['label']}。")
        append_monitor_run(session, task.id, "login", "info", f"登录入口：{platform_meta['login_url']}")
        results = []
        search_provider = session.scalar(select(SearchProvider).where(SearchProvider.enabled == 1).order_by(SearchProvider.id.asc()))
        try:
            if search_provider:
                results = search_with_tavily(search_provider, task.keyword)
                append_monitor_run(session, task.id, "candidate_search", "info", f"已保存 {len(results)} 条候选搜索结果，仅用于人工对照，不计入官方引用统计。")
            else:
                append_monitor_run(session, task.id, "candidate_search", "warning", "未配置搜索通道，跳过候选信源池；不影响官方网页抓取。")
        except HTTPException as error:
            append_monitor_run(session, task.id, "candidate_search", "warning", f"候选搜索失败：{error.detail}；继续执行官方网页抓取。")
        try:
            worker_result = run_official_monitor_worker(task.platform, task.keyword, task.task_code)
            append_monitor_run(session, task.id, "browser_worker", worker_result.get("status", "info"), worker_result.get("message", "浏览器 Worker 已执行"))
        except Exception as error:
            worker_result = {"status": "failed", "message": str(error), "answer_text": "", "sources": [], "screenshot_path": ""}
            append_monitor_run(session, task.id, "browser_worker", "error", str(error))
        captured_sources = worker_result.get("sources", []) if isinstance(worker_result.get("sources", []), list) else []
        answer_text = str(worker_result.get("answer_text") or "").strip()
        if worker_result.get("status") == "login_required":
            task.status = "login_required"
            task.error_message = f"{platform_meta['label']} 当前未登录，请先在浏览器完成登录后重试。"
            answer_text = answer_text or f"{platform_meta['label']} 尚未登录，当前暂存候选信源池供人工对照。"
        elif worker_result.get("status") == "blocked_by_challenge":
            task.status = "blocked_by_challenge"
            task.error_message = f"{platform_meta['label']} 当前被平台风控或人机验证拦截，请使用人工登录后的真实 Chrome 会话重试。"
            answer_text = answer_text or task.error_message
        elif worker_result.get("status") == "completed":
            task.status = "completed"
            task.error_message = ""
        else:
            task.status = "completed" if captured_sources else "failed"
            task.error_message = "" if captured_sources else str(worker_result.get("message") or "浏览器 Worker 未返回可用结果")
        if not answer_text:
            answer_text = f"当前展示的是 {platform_meta['label']} 对关键词“{task.keyword}”的候选信源池。待真实网页登录成功后，系统会替换为网页端显式引用。"
        session.add(OfficialMonitorAnswer(
            task_id=task.id,
            answer_text=answer_text,
            answer_html="",
            raw_capture_json=json.dumps({"platform": task.platform, "keyword": task.keyword, "worker_result": worker_result, "candidate_results": results}, ensure_ascii=False),
            screenshot_path=str(worker_result.get("screenshot_path") or ""),
        ))
        source_rows = captured_sources
        for index, item in enumerate(source_rows, start=1):
            item_url = str(item.get("url") or "").strip()
            if not item_url:
                continue
            title = str(item.get("title") or item_url)
            display_text = str(item.get("snippet") or item.get("display_text") or "")
            normalized_url = str(item.get("normalized_url") or item_url)
            host, root = normalize_domain(item_url)
            session.add(OfficialMonitorSource(
                task_id=task.id,
                title=title,
                url=item_url,
                normalized_url=normalized_url,
                domain=host,
                root_domain=root,
                source_type=str(item.get("source_type") or source_type_for_domain(host)),
                rank_index=int(item.get("rank_index") or index),
                display_text=display_text,
            ))
        append_monitor_run(session, task.id, "capture", "info", "已保存浏览器结果、候选信源池和运行日志。")
        task.finished_at = datetime.now()
        log(session, "官方信源监测", f"{platform_meta['label']} · {task.keyword}", "查询")
        session.commit()


def provider_cipher() -> Fernet:
    configured_key = os.getenv("GEO_CONFIG_ENCRYPTION_KEY")
    key_path = Path(__file__).resolve().parent.parent / ".geo_config.key"
    if configured_key:
        key = configured_key.encode()
    elif key_path.exists():
        key = key_path.read_bytes().strip()
    else:
        key = Fernet.generate_key()
        key_path.write_bytes(key)
        key_path.chmod(0o600)
    return Fernet(key)


def serialize_article(article: GeneratedArticle):
    result = serialize(article)
    result["include_faq"] = bool(article.include_faq)
    result["include_knowledge"] = bool(article.include_knowledge)
    try:
        result["selected_images"] = json.loads(article.selected_images or "[]")
    except json.JSONDecodeError:
        result["selected_images"] = []
    return result


def serialize_geo_report(report: GeoReport):
    return {
        "id": report.id,
        "report_code": report.report_code,
        "keyword": report.keyword,
        "provider_names": json.loads(report.provider_names or "[]"),
        "created_at": report.created_at.isoformat(),
    }


def geo_report_payload(report: GeoReport):
    try:
        payload = json.loads(report.result_json)
    except json.JSONDecodeError as error:
        raise HTTPException(500, "报告快照数据损坏，无法读取") from error
    payload["aggregate"] = aggregate_geo_results(payload)
    payload["report"] = serialize_geo_report(report)
    return payload


def geo_report_or_404(session: Session, report_id: int):
    report = session.get(GeoReport, report_id)
    if not report:
        raise HTTPException(404, "评估报告不存在")
    return report


def report_download_headers(report: GeoReport, extension: str):
    safe_keyword = "".join(character for character in report.keyword if character not in '\\/:*?"<>|').strip() or "GEO-report"
    filename = f"{safe_keyword}-{report.report_code}.{extension}"
    ascii_filename = f"{report.report_code}.{extension}"
    return {"Content-Disposition": f"attachment; filename=\"{ascii_filename}\"; filename*=UTF-8''{quote(filename)}"}


def aggregate_geo_results(payload: dict) -> dict:
    domain_rows: dict[str, dict] = {}
    page_rows: dict[str, dict] = {}
    results = payload.get("results", {}) if isinstance(payload.get("results"), dict) else {}
    for platform, result in results.items():
        if not isinstance(result, dict):
            continue
        for page in result.get("pages", []):
            if not isinstance(page, dict) or not page.get("url"):
                continue
            domain = str(page.get("domain") or urlparse(str(page["url"])).netloc).lower()
            confirmed = page.get("mentions")
            confirmed_count = int(confirmed) if isinstance(confirmed, int) and confirmed > 0 else 0
            source_record = 1
            domain_entry = domain_rows.setdefault(domain, {
                "domain": domain,
                "site_name": str(page.get("site_name") or domain),
                "confirmed_mentions": 0,
                "source_records": 0,
                "urls": set(),
                "models": {},
            })
            if domain_entry["site_name"] == domain and page.get("site_name"):
                domain_entry["site_name"] = str(page["site_name"])
            domain_entry["confirmed_mentions"] += confirmed_count
            domain_entry["source_records"] += source_record
            domain_entry["urls"].add(str(page["url"]))
            model_entry = domain_entry["models"].setdefault(platform, {"confirmed_mentions": 0, "source_records": 0})
            model_entry["confirmed_mentions"] += confirmed_count
            model_entry["source_records"] += source_record

            page_entry = page_rows.setdefault(str(page["url"]), {
                "title": str(page.get("title") or page["url"]),
                "url": str(page["url"]),
                "domain": domain,
                "site_name": str(page.get("site_name") or domain),
                "date": str(page.get("date") or "未知"),
                "confirmed_mentions": 0,
                "source_records": 0,
                "models": {},
            })
            page_entry["confirmed_mentions"] += confirmed_count
            page_entry["source_records"] += source_record
            page_model = page_entry["models"].setdefault(platform, {"confirmed_mentions": 0, "source_records": 0})
            page_model["confirmed_mentions"] += confirmed_count
            page_model["source_records"] += source_record

    total_confirmed = sum(row["confirmed_mentions"] for row in domain_rows.values())
    domains = []
    for row in domain_rows.values():
        models = [
            {"name": name, **values}
            for name, values in sorted(row["models"].items())
        ]
        domains.append({
            "domain": row["domain"],
            "site_name": row["site_name"],
            "confirmed_mentions": row["confirmed_mentions"],
            "source_records": row["source_records"],
            "model_count": len(models),
            "page_count": len(row["urls"]),
            "share": f"{round(row['confirmed_mentions'] * 100 / total_confirmed)}%" if total_confirmed else "未提供",
            "models": models,
        })
    pages = []
    for row in page_rows.values():
        models = [
            {"name": name, **values}
            for name, values in sorted(row["models"].items())
        ]
        pages.append({
            "title": row["title"], "url": row["url"], "domain": row["domain"], "site_name": row["site_name"], "date": row["date"],
            "confirmed_mentions": row["confirmed_mentions"], "source_records": row["source_records"], "model_count": len(models), "models": models,
        })
    domains.sort(key=lambda row: (-row["confirmed_mentions"], -row["model_count"], -row["source_records"], row["domain"]))
    pages.sort(key=lambda row: (-row["confirmed_mentions"], -row["model_count"], -row["source_records"], row["url"]))
    return {
        "total_confirmed_mentions": total_confirmed,
        "total_source_records": sum(row["source_records"] for row in domain_rows.values()),
        "domain_count": len(domains),
        "page_count": len(pages),
        "domains": domains,
        "pages": pages,
    }


def set_cell_shading(cell, fill: str):
    cell_properties = cell._tc.get_or_add_tcPr()
    shading = cell_properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        cell_properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int):
    cell_properties = cell._tc.get_or_add_tcPr()
    cell_width = cell_properties.find(qn("w:tcW"))
    if cell_width is None:
        cell_width = OxmlElement("w:tcW")
        cell_properties.append(cell_width)
    cell_width.set(qn("w:w"), str(width_dxa))
    cell_width.set(qn("w:type"), "dxa")


def set_docx_run_font(run, size: float = 10.5, bold: bool = False, color: str = "344054"):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_docx_text(doc: Document, text_value: str, size: float = 10.5, bold: bool = False, color: str = "344054", after: float = 6, alignment=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.1
    if alignment is not None:
        paragraph.alignment = alignment
    set_docx_run_font(paragraph.add_run(text_value), size=size, bold=bold, color=color)
    return paragraph


def add_docx_heading(doc: Document, text_value: str, level: int = 1):
    sizes = {1: 16, 2: 13, 3: 11.5}
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    paragraph.paragraph_format.space_after = Pt(7 if level == 1 else 5)
    paragraph.paragraph_format.keep_with_next = True
    set_docx_run_font(paragraph.add_run(text_value), size=sizes[level], bold=True, color="2E74B5" if level < 3 else "1F4D78")
    return paragraph


def style_docx_table(table, widths: list[int], header: bool = True):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table_properties = table._tbl.tblPr
    table_width = table_properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_properties.append(table_width)
    table_width.set(qn("w:w"), str(sum(widths)))
    table_width.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row_index, row in enumerate(table.rows):
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if header and row_index == 0:
                set_cell_shading(cell, "F2F4F7")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    set_docx_run_font(run, size=9.2, bold=header and row_index == 0, color="344054")


MARKDOWN_PARSER = MarkdownIt("commonmark").enable("table")


def markdown_tree(text_value: str):
    return SyntaxTreeNode(MARKDOWN_PARSER.parse(text_value or ""))


def markdown_plain_text(node) -> str:
    if node.type in {"text", "code_inline", "code_block", "fence"}:
        return node.content
    if node.type in {"softbreak", "hardbreak"}:
        return "\n"
    return "".join(markdown_plain_text(child) for child in (node.children or []))


def add_docx_inline(paragraph, node, bold: bool = False, italic: bool = False):
    if node.type in {"text", "code_inline"}:
        run = paragraph.add_run(node.content)
        set_docx_run_font(run, size=10.5, bold=bold)
        run.italic = italic
        if node.type == "code_inline":
            run.font.name = "Courier New"
        return
    if node.type in {"softbreak", "hardbreak"}:
        paragraph.add_run().add_break()
        return
    next_bold = bold or node.type == "strong"
    next_italic = italic or node.type == "em"
    for child in node.children or []:
        add_docx_inline(paragraph, child, next_bold, next_italic)


def add_docx_markdown(document: Document, text_value: str):
    for node in markdown_tree(text_value).children or []:
        if node.type == "heading":
            level = int((node.tag or "h2")[1:]) if (node.tag or "h2")[1:].isdigit() else 2
            add_docx_heading(document, markdown_plain_text(node).strip(), min(max(level, 1), 3))
        elif node.type == "paragraph":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.1
            for child in node.children or []:
                add_docx_inline(paragraph, child)
        elif node.type in {"bullet_list", "ordered_list"}:
            style_name = "List Bullet" if node.type == "bullet_list" else "List Number"
            for item in node.children or []:
                paragraph = document.add_paragraph(style=style_name)
                paragraph.paragraph_format.space_after = Pt(4)
                for child in item.children or []:
                    for inline in child.children or []:
                        add_docx_inline(paragraph, inline)
        elif node.type == "blockquote":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.2)
            paragraph.paragraph_format.space_after = Pt(6)
            set_docx_run_font(paragraph.add_run(markdown_plain_text(node).strip()), size=10.5, color="475467")
        elif node.type == "table":
            rows = []
            for group in node.children or []:
                for row in group.children or []:
                    rows.append([markdown_plain_text(cell).strip() for cell in row.children or []])
            if not rows:
                continue
            column_count = max(len(row) for row in rows)
            table = document.add_table(rows=1, cols=column_count)
            for row_index, values in enumerate(rows):
                cells = table.rows[0].cells if row_index == 0 else table.add_row().cells
                for index, value in enumerate(values):
                    cells[index].text = value
            style_docx_table(table, [9360 // column_count] * column_count)
        elif node.type in {"fence", "code_block"}:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(6)
            run = paragraph.add_run(node.content)
            run.font.name = "Courier New"
            run.font.size = Pt(9)


def display_mentions(value) -> str:
    return f"{value} 次" if isinstance(value, int) else "未提供"


def build_geo_report_docx(report: GeoReport) -> bytes:
    payload = json.loads(report.result_json)
    aggregate = aggregate_geo_results(payload)
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    header = section.header.paragraphs[0]
    set_docx_run_font(header.add_run("GEO 智能优化引擎 | 评估报告"), size=8.5, color="667085")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_docx_run_font(footer.add_run(report.report_code), size=8.5, color="98A2B3")

    add_docx_text(document, "GEO 评估报告", size=23, bold=True, color="1F2937", after=4)
    add_docx_text(document, f"关键词：{report.keyword}", size=13, bold=True, color="475467", after=3)
    add_docx_text(document, f"报告编号：{report.report_code}    生成时间：{report.created_at.strftime('%Y-%m-%d %H:%M')}", size=9.5, color="667085", after=16)

    metrics = [("参与模型", len(payload["platforms"])), ("确认引用", aggregate["total_confirmed_mentions"]), ("来源记录", aggregate["total_source_records"]), ("高频网站", aggregate["domain_count"])]
    metric_table = document.add_table(rows=2, cols=4)
    for index, (label, value) in enumerate(metrics):
        metric_table.cell(0, index).text = str(label)
        metric_table.cell(1, index).text = str(value)
        set_cell_shading(metric_table.cell(0, index), "F2F4F7")
        metric_table.cell(0, index).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        metric_table.cell(1, index).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_docx_table(metric_table, [2340, 2340, 2340, 2340], header=False)
    for index in range(4):
        for run in metric_table.cell(0, index).paragraphs[0].runs:
            set_docx_run_font(run, size=8.8, bold=True, color="667085")
        for run in metric_table.cell(1, index).paragraphs[0].runs:
            set_docx_run_font(run, size=16, bold=True, color="2563EB")

    add_docx_heading(document, "1. 评估汇总", 1)
    add_docx_text(document, f"本次调用 {len(payload['platforms'])} 个已启用模型：{'、'.join(payload['platforms'])}。共保留 {aggregate['total_source_records']} 条来源记录，其中可确认正文引用 {aggregate['total_confirmed_mentions']} 次。以下排行按确认引用次数、模型覆盖和来源记录依次排序，用于筛选优先投放内容的网站。", after=8)
    add_docx_heading(document, "高频引用网站排行", 2)
    summary_table = document.add_table(rows=1, cols=5)
    for index, title in enumerate(["网站 / 域名", "确认引用", "模型覆盖", "来源记录", "页面数"]):
        summary_table.cell(0, index).text = title
    for domain in aggregate["domains"]:
        cells = summary_table.add_row().cells
        cells[0].text = domain["site_name"] or domain["domain"]
        cells[1].text = str(domain["confirmed_mentions"])
        cells[2].text = f"{domain['model_count']} 个模型"
        cells[3].text = str(domain["source_records"])
        cells[4].text = str(domain["page_count"])
    style_docx_table(summary_table, [3000, 1350, 1700, 1700, 1610])
    add_docx_heading(document, "跨模型页面累计明细", 2)
    page_summary_table = document.add_table(rows=1, cols=5)
    for index, title in enumerate(["页面标题", "域名", "确认引用", "模型覆盖", "URL"]):
        page_summary_table.cell(0, index).text = title
    for page in aggregate["pages"]:
        cells = page_summary_table.add_row().cells
        cells[0].text = page["title"]
        cells[1].text = page["domain"]
        cells[2].text = str(page["confirmed_mentions"])
        cells[3].text = f"{page['model_count']} 个模型"
        cells[4].text = page["url"]
    style_docx_table(page_summary_table, [2150, 1500, 1200, 1500, 3010])

    for model_index, platform in enumerate(payload["platforms"], start=1):
        item = payload["results"][platform]
        add_docx_heading(document, f"{model_index}. {platform}", 1)
        add_docx_text(document, item.get("model", platform), size=9.5, bold=True, color="667085", after=5)
        add_docx_heading(document, "模型回答", 2)
        add_docx_markdown(document, item.get("summary", ""))
        add_docx_heading(document, "引用域名", 2)
        domain_table = document.add_table(rows=1, cols=3)
        for index, title in enumerate(["域名", "正文引用次数", "占比"]):
            domain_table.cell(0, index).text = title
        for domain in item.get("domains", []):
            cells = domain_table.add_row().cells
            cells[0].text = str(domain.get("domain", ""))
            cells[1].text = display_mentions(domain.get("mentions"))
            cells[2].text = str(domain.get("share", ""))
        style_docx_table(domain_table, [5900, 1730, 1730])
        add_docx_heading(document, "引用页面", 2)
        page_table = document.add_table(rows=1, cols=5)
        for index, title in enumerate(["页面标题", "站点", "正文引用次数", "URL", "日期"]):
            page_table.cell(0, index).text = title
        for page in item.get("pages", []):
            cells = page_table.add_row().cells
            cells[0].text = page.get("title") or page.get("url", "")
            cells[1].text = page.get("site_name") or page.get("domain", "")
            cells[2].text = display_mentions(page.get("mentions"))
            cells[3].text = page.get("url", "")
            cells[4].text = page.get("date", "未知")
        style_docx_table(page_table, [2000, 1300, 1300, 3500, 1260])

    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def pdf_font_name() -> str:
    name = "GeoCJK"
    if name not in pdfmetrics.getRegisteredFontNames():
        font_path = "/System/Library/Fonts/Supplemental/Arial Unicode.ttf"
        pdfmetrics.registerFont(TTFont(name, font_path))
    return name


def pdf_inline_markup(node) -> str:
    if node.type == "text":
        return escape(node.content, quote=False)
    if node.type == "code_inline":
        return f"<font name=\"Courier\">{escape(node.content, quote=False)}</font>"
    if node.type in {"softbreak", "hardbreak"}:
        return "<br/>"
    content = "".join(pdf_inline_markup(child) for child in (node.children or []))
    if node.type == "strong":
        return f"<b>{content}</b>"
    if node.type == "em":
        return f"<i>{content}</i>"
    if node.type == "link":
        return f"<u>{content}</u>"
    return content


def pdf_markdown_flowables(text_value: str, body_style, h1_style, h2_style, cell_style, header_style):
    flowables = []
    for node in markdown_tree(text_value).children or []:
        if node.type == "heading":
            level = int((node.tag or "h2")[1:]) if (node.tag or "h2")[1:].isdigit() else 2
            style = h1_style if level <= 2 else h2_style
            flowables.append(Paragraph(pdf_inline_markup(node), style))
        elif node.type == "paragraph":
            flowables.append(Paragraph(pdf_inline_markup(node), body_style))
        elif node.type in {"bullet_list", "ordered_list"}:
            for number, item in enumerate(node.children or [], start=1):
                prefix = "•" if node.type == "bullet_list" else f"{number}."
                item_text = " ".join(pdf_inline_markup(child) for child in item.children or [])
                flowables.append(Paragraph(f"{prefix} {item_text}", body_style))
        elif node.type == "blockquote":
            quote_style = ParagraphStyle("GeoQuote", parent=body_style, leftIndent=12, borderColor=colors.HexColor("#93C5FD"), borderWidth=2, borderPadding=7, backColor=colors.HexColor("#F8FBFF"))
            flowables.append(Paragraph(pdf_inline_markup(node), quote_style))
        elif node.type == "table":
            rows = []
            for group in node.children or []:
                for row in group.children or []:
                    rows.append([pdf_inline_markup(cell) for cell in row.children or []])
            if rows:
                col_count = max(len(row) for row in rows)
                data = []
                for row_index, row in enumerate(rows):
                    style = header_style if row_index == 0 else cell_style
                    data.append([Paragraph(value, style) for value in row + [""] * (col_count - len(row))])
                table = Table(data, colWidths=[6.5 * inch / col_count] * col_count, repeatRows=1)
                table.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")), ("GRID", (0, 0), (-1, -1), 0.45, colors.HexColor("#D0D5DD")), ("VALIGN", (0, 0), (-1, -1), "TOP"), ("TOPPADDING", (0, 0), (-1, -1), 5), ("BOTTOMPADDING", (0, 0), (-1, -1), 5), ("LEFTPADDING", (0, 0), (-1, -1), 5), ("RIGHTPADDING", (0, 0), (-1, -1), 5)]))
                flowables.append(table)
        elif node.type in {"fence", "code_block"}:
            code_style = ParagraphStyle("GeoCode", parent=cell_style, backColor=colors.HexColor("#F8FAFC"), borderColor=colors.HexColor("#E4E7EC"), borderWidth=0.5, borderPadding=7)
            flowables.append(Paragraph(escape(node.content, quote=False).replace("\n", "<br/>"), code_style))
    return flowables


def build_geo_report_pdf(report: GeoReport) -> bytes:
    payload = json.loads(report.result_json)
    aggregate = aggregate_geo_results(payload)
    output = BytesIO()
    font = pdf_font_name()
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("GeoTitle", parent=styles["Title"], fontName=font, fontSize=23, leading=28, textColor=colors.HexColor("#1F2937"), spaceAfter=8)
    subtitle_style = ParagraphStyle("GeoSubtitle", parent=styles["Normal"], fontName=font, fontSize=12, leading=16, textColor=colors.HexColor("#475467"), spaceAfter=4)
    body_style = ParagraphStyle("GeoBody", parent=styles["BodyText"], fontName=font, fontSize=9.5, leading=14, textColor=colors.HexColor("#344054"), spaceAfter=7)
    h1_style = ParagraphStyle("GeoH1", parent=styles["Heading1"], fontName=font, fontSize=15, leading=19, textColor=colors.HexColor("#2E74B5"), spaceBefore=12, spaceAfter=7)
    h2_style = ParagraphStyle("GeoH2", parent=styles["Heading2"], fontName=font, fontSize=11.5, leading=15, textColor=colors.HexColor("#2E74B5"), spaceBefore=9, spaceAfter=5)
    cell_style = ParagraphStyle("GeoCell", parent=body_style, fontSize=7.8, leading=10, spaceAfter=0)
    header_style = ParagraphStyle("GeoHeader", parent=cell_style, fontSize=8, leading=10, textColor=colors.HexColor("#344054"), alignment=TA_CENTER)
    document = SimpleDocTemplate(output, pagesize=letter, rightMargin=inch, leftMargin=inch, topMargin=inch, bottomMargin=inch, title=f"GEO 评估报告 - {report.keyword}")
    pdf_text = lambda value: escape(str(value), quote=False)
    story = [Paragraph("GEO 评估报告", title_style), Paragraph(f"关键词：{pdf_text(report.keyword)}", subtitle_style), Paragraph(f"报告编号：{pdf_text(report.report_code)}　生成时间：{report.created_at.strftime('%Y-%m-%d %H:%M')}", body_style), Spacer(1, 8)]
    metrics = [[Paragraph(label, header_style) for label in ["参与模型", "确认引用", "来源记录", "高频网站"]], [Paragraph(str(value), ParagraphStyle("Metric", parent=header_style, fontSize=15, leading=19, textColor=colors.HexColor("#2563EB"))) for value in [len(payload["platforms"]), aggregate["total_confirmed_mentions"], aggregate["total_source_records"], aggregate["domain_count"]]]]
    metric_table = Table(metrics, colWidths=[1.625 * inch] * 4)
    metric_table.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")), ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#D0D5DD")), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("TOPPADDING", (0, 0), (-1, -1), 7), ("BOTTOMPADDING", (0, 0), (-1, -1), 7)]))
    story.extend([metric_table, Paragraph("1. 评估汇总", h1_style), Paragraph(pdf_text(f"本次调用 {len(payload['platforms'])} 个已启用模型：{'、'.join(payload['platforms'])}。共保留 {aggregate['total_source_records']} 条来源记录，其中可确认正文引用 {aggregate['total_confirmed_mentions']} 次。以下排行按确认引用次数、模型覆盖和来源记录依次排序，用于筛选优先投放内容的网站。"), body_style)])

    def source_table(pages: list[dict], include_mentions: bool = True):
        headers = ["页面标题", "站点", "正文引用次数", "URL", "日期"] if include_mentions else ["页面标题", "站点", "URL", "日期"]
        rows = [[Paragraph(value, header_style) for value in headers]]
        for page in pages:
            values = [Paragraph(pdf_text(page.get("title") or page.get("url", "")), cell_style), Paragraph(pdf_text(page.get("site_name") or page.get("domain", "")), cell_style)]
            if include_mentions:
                values.append(Paragraph(pdf_text(display_mentions(page.get("mentions"))), cell_style))
            values.extend([Paragraph(pdf_text(page.get("url", "")), cell_style), Paragraph(pdf_text(page.get("date", "未知")), cell_style)])
            rows.append(values)
        widths = [1.3 * inch, .9 * inch, .75 * inch, 2.65 * inch, .9 * inch] if include_mentions else [1.75 * inch, 1.05 * inch, 2.8 * inch, .9 * inch]
        table = Table(rows, colWidths=widths, repeatRows=1)
        table.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")), ("GRID", (0, 0), (-1, -1), 0.45, colors.HexColor("#D0D5DD")), ("VALIGN", (0, 0), (-1, -1), "TOP"), ("LEFTPADDING", (0, 0), (-1, -1), 5), ("RIGHTPADDING", (0, 0), (-1, -1), 5), ("TOPPADDING", (0, 0), (-1, -1), 5), ("BOTTOMPADDING", (0, 0), (-1, -1), 5)]))
        return table

    story.append(Paragraph("高频引用网站排行", h2_style))
    domain_rows = [[Paragraph(value, header_style) for value in ["网站 / 域名", "确认引用", "模型覆盖", "来源记录", "页面数"]]]
    for domain in aggregate["domains"]:
        domain_rows.append([Paragraph(pdf_text(domain["site_name"] or domain["domain"]), cell_style), Paragraph(str(domain["confirmed_mentions"]), cell_style), Paragraph(f"{domain['model_count']} 个模型", cell_style), Paragraph(str(domain["source_records"]), cell_style), Paragraph(str(domain["page_count"]), cell_style)])
    aggregate_domains = Table(domain_rows, colWidths=[2.1 * inch, .85 * inch, 1.1 * inch, 1.1 * inch, 1.35 * inch], repeatRows=1)
    aggregate_domains.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")), ("GRID", (0, 0), (-1, -1), 0.45, colors.HexColor("#D0D5DD")), ("VALIGN", (0, 0), (-1, -1), "TOP"), ("TOPPADDING", (0, 0), (-1, -1), 5), ("BOTTOMPADDING", (0, 0), (-1, -1), 5)]))
    story.extend([aggregate_domains, Paragraph("跨模型页面累计明细", h2_style)])
    aggregate_pages = [{"title": page["title"], "site_name": page["domain"], "mentions": page["confirmed_mentions"], "url": page["url"], "date": page["date"]} for page in aggregate["pages"]]
    story.append(source_table(aggregate_pages))
    for model_index, platform in enumerate(payload["platforms"], start=1):
        item = payload["results"][platform]
        story.extend([PageBreak(), Paragraph(pdf_text(f"{model_index}. {platform}"), h1_style), Paragraph(pdf_text(item.get("model", platform)), body_style), Paragraph("模型回答", h2_style)])
        story.extend(pdf_markdown_flowables(item.get("summary", ""), body_style, h1_style, h2_style, cell_style, header_style))
        story.append(Paragraph("引用域名", h2_style))
        domain_rows = [[Paragraph(value, header_style) for value in ["域名", "正文引用次数", "占比"]]] + [[Paragraph(pdf_text(row.get("domain", "")), cell_style), Paragraph(pdf_text(display_mentions(row.get("mentions"))), cell_style), Paragraph(pdf_text(row.get("share", "")), cell_style)] for row in item.get("domains", [])]
        domain_table = Table(domain_rows, colWidths=[4.1 * inch, 1.2 * inch, 1.2 * inch], repeatRows=1)
        domain_table.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")), ("GRID", (0, 0), (-1, -1), 0.45, colors.HexColor("#D0D5DD")), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("TOPPADDING", (0, 0), (-1, -1), 5), ("BOTTOMPADDING", (0, 0), (-1, -1), 5)]))
        story.extend([domain_table, Paragraph("引用页面", h2_style), source_table(item.get("pages", []))])

    def page_footer(canvas, doc):
        canvas.saveState()
        canvas.setFont(font, 8)
        canvas.setFillColor(colors.HexColor("#98A2B3"))
        canvas.drawString(inch, 0.5 * inch, "GEO 智能优化引擎 | 评估报告")
        canvas.drawRightString(7.5 * inch, 0.5 * inch, f"{report.report_code}  |  {doc.page}")
        canvas.restoreState()

    document.build(story, onFirstPage=page_footer, onLaterPages=page_footer)
    return output.getvalue()


def log(session: Session, action: str, detail: str = "", category: str | None = None):
    inferred = "删除" if "删除" in action else "修改" if any(word in action for word in ["编辑", "保存", "修改"]) else "生成" if any(word in action for word in ["生成", "成文", "确认", "翻译"]) else "查询" if any(word in action for word in ["查询", "搜索"]) else "新增"
    session.add(AuditLog(category=category or inferred, action=action, detail=detail))


def seed():
    Base.metadata.create_all(engine)
    with SessionLocal() as session:
        media_columns = {row[1] for row in session.execute(text("PRAGMA table_info(media_providers)"))}
        if not media_columns:
            # create_all above creates the table; this branch is retained for old SQLite files.
            Base.metadata.create_all(engine)
        article_columns = {row[1] for row in session.execute(text("PRAGMA table_info(generated_articles)"))}
        if "content_saved_at" not in article_columns:
            session.execute(text("ALTER TABLE generated_articles ADD COLUMN content_saved_at DATETIME"))
            session.commit()
        provider_columns = {row[1] for row in session.execute(text("PRAGMA table_info(model_providers)"))}
        if "provider_type" not in provider_columns:
            session.execute(text("ALTER TABLE model_providers ADD COLUMN provider_type TEXT DEFAULT 'qwen'"))
            session.commit()
        if "api_secret_encrypted" not in provider_columns:
            session.execute(text("ALTER TABLE model_providers ADD COLUMN api_secret_encrypted TEXT DEFAULT ''"))
            session.commit()
        if "aux_api_key_encrypted" not in provider_columns:
            session.execute(text("ALTER TABLE model_providers ADD COLUMN aux_api_key_encrypted TEXT DEFAULT ''"))
            session.commit()
        columns = {row[1] for row in session.execute(text("PRAGMA table_info(brands)"))}
        migrations = {"brand_code": "TEXT DEFAULT ''", "founded": "TEXT DEFAULT ''", "products": "TEXT DEFAULT ''", "target_customers": "TEXT DEFAULT ''", "core_capability": "TEXT DEFAULT ''", "certifications": "TEXT DEFAULT ''"}
        for column, definition in migrations.items():
            if column not in columns:
                session.execute(text(f"ALTER TABLE brands ADD COLUMN {column} {definition}"))
        session.commit()
        faq_columns = {row[1] for row in session.execute(text("PRAGMA table_info(faqs)"))}
        if "faq_code" not in faq_columns:
            session.execute(text("ALTER TABLE faqs ADD COLUMN faq_code TEXT DEFAULT ''"))
            session.commit()
        log_columns = {row[1] for row in session.execute(text("PRAGMA table_info(audit_logs)"))}
        if "category" not in log_columns:
            session.execute(text("ALTER TABLE audit_logs ADD COLUMN category TEXT DEFAULT '其他'"))
            session.commit()
        for entry in session.scalars(select(AuditLog).where(AuditLog.category == "其他")).all():
            entry.category = "删除" if "删除" in entry.action else "修改" if any(word in entry.action for word in ["编辑", "保存", "修改", "更新"]) else "生成" if any(word in entry.action for word in ["生成", "成文", "确认", "翻译"]) else "查询" if any(word in entry.action for word in ["查询", "搜索"]) else "新增"
        session.commit()
        if not session.scalar(select(func.count()).select_from(Brand)):
            session.add_all([
            Brand(brand_code="BRD-001", name="品牌A", company="品牌A科技有限公司", alias="BrandA", founded="2018", region="上海", products="SaaS 工具", target_customers="中小企业", core_capability="低代码平台", certifications="ISO 27001", constraint=DEFAULT_CONSTRAINT),
            Brand(brand_code="BRD-002", name="品牌B", company="品牌B数字营销有限公司", alias="BrandB", founded="2020", region="北京", products="数字营销", target_customers="大型企业", core_capability="AI 投放优化", certifications="CMMI 3", constraint=DEFAULT_CONSTRAINT),
            Faq(faq_code="FAQ-001", category="采购流程", question="如何通过 Zeta Voltage 采购设备？", answer="您可以选择产品并申请合并报价，或提交询价单。工程师将在 1 个工作日内响应，并在确认技术参数后提供完整方案与交期。"),
            Faq(faq_code="FAQ-002", category="产品描述与认证", question="产品是否符合 SNI、IEC、CE 和 ISO 标准？", answer="是的。我们可按项目和目标市场提供适用的标准、认证与检验资料，协助完成采购前的合规核验。"),
            Faq(faq_code="FAQ-003", category="价格与起订量", question="最低起订量（MOQ）是多少？", answer="不同设备型号的 MOQ 和报价条件不同。支持样品、试订单和批量采购等灵活方式，具体以技术规格和订单数量为准。"),
            Faq(faq_code="FAQ-004", category="售后与质保", question="产品质保期和售后响应如何安排？", answer="我们提供清晰的质保范围、交付资料和技术支持路径。出现运行问题时，服务团队将根据故障信息提供远程或项目级处理建议。"),
            Knowledge(name="VIP 贵宾厅服务", category="产品与服务", content="覆盖全国 50 余个机场，提供私密休息区、专属餐饮、快速安检通道等服务。"),
            Knowledge(name="500 强差旅合作案例", category="客户案例", content="为客户定制差旅方案，年差旅成本下降 23%，员工满意度提升 35%。"),
        ])
            log(session, "系统初始化", "已创建演示运营数据")
            session.commit()
        else:
            # Existing v2.3 demo records are normalized once to the requested baseline data.
            existing = session.scalars(select(Brand).order_by(Brand.id)).all()
            if len(existing) == 2 and {item.name for item in existing} <= {"海云端", "品牌 B"}:
                baseline = [("BRD-001", "品牌A", "品牌A科技有限公司", "BrandA", "2018", "上海", "SaaS 工具", "中小企业", "低代码平台", "ISO 27001"), ("BRD-002", "品牌B", "品牌B数字营销有限公司", "BrandB", "2020", "北京", "数字营销", "大型企业", "AI 投放优化", "CMMI 3")]
                for item, values in zip(existing, baseline):
                    item.brand_code, item.name, item.company, item.alias, item.founded, item.region, item.products, item.target_customers, item.core_capability, item.certifications = values
                    item.constraint = DEFAULT_CONSTRAINT
                session.commit()
        existing_faqs = session.scalars(select(Faq).order_by(Faq.id)).all()
        legacy_questions = {"如何开始使用海云端商旅服务？", "服务覆盖哪些机场？", "行程变更如何处理？"}
        if existing_faqs and {faq.question for faq in existing_faqs} <= legacy_questions:
            session.query(Faq).delete()
            session.add_all([
                Faq(faq_code="FAQ-001", category="采购流程", question="如何通过 Zeta Voltage 采购设备？", answer="您可以选择产品并申请合并报价，或提交询价单。工程师将在 1 个工作日内响应，并在确认技术参数后提供完整方案与交期。"),
                Faq(faq_code="FAQ-002", category="产品描述与认证", question="产品是否符合 SNI、IEC、CE 和 ISO 标准？", answer="是的。我们可按项目和目标市场提供适用的标准、认证与检验资料，协助完成采购前的合规核验。"),
                Faq(faq_code="FAQ-003", category="价格与起订量", question="最低起订量（MOQ）是多少？", answer="不同设备型号的 MOQ 和报价条件不同。支持样品、试订单和批量采购等灵活方式，具体以技术规格和订单数量为准。"),
                Faq(faq_code="FAQ-004", category="售后与质保", question="产品质保期和售后响应如何安排？", answer="我们提供清晰的质保范围、交付资料和技术支持路径。出现运行问题时，服务团队将根据故障信息提供远程或项目级处理建议。"),
            ])
            session.commit()
        for faq in session.scalars(select(Faq).where(Faq.faq_code == "")).all():
            faq.faq_code = f"FAQ-{faq.id:03d}"
        for legacy in session.scalars(select(ModelProvider).where(ModelProvider.name.in_(["GPT-5.5", "DeepSeek"]))).all():
            legacy.provider_type = "legacy"
            legacy.enabled = 0
        for hunyuan_provider in session.scalars(select(ModelProvider).where(ModelProvider.provider_type == "hunyuan")).all():
            # Migrate incomplete legacy records to the TokenHub OpenAI-compatible API.
            if hunyuan_provider.base_url.rstrip("/") == "https://hunyuan.tencentcloudapi.com" and not hunyuan_provider.api_secret_encrypted:
                hunyuan_provider.base_url = "https://tokenhub-intl.tencentmaas.com/v1"
            hunyuan_provider.protocol = "chat_completions"
        for provider in session.scalars(select(ModelProvider).where(ModelProvider.provider_type.in_(["qwen", "hunyuan", "volcengine", "tencent_search"]))).all():
            uses_legacy_hunyuan_auth = provider.provider_type == "hunyuan" and "tencentmaas.com" not in provider.base_url.lower()
            uses_tencent_search = provider.provider_type == "tencent_search"
            credentials_ready = bool(provider.api_key_encrypted) and (not uses_legacy_hunyuan_auth or bool(provider.api_secret_encrypted)) and (not uses_tencent_search or bool(provider.aux_api_key_encrypted))
            if not credentials_ready:
                provider.enabled = 0
        volcengine = session.scalar(select(ModelProvider).where(ModelProvider.provider_type == "volcengine").order_by(ModelProvider.id.asc()))
        media_defaults = [
            ("豆包配图 · Seedream 5.0 Lite", "image", "doubao-seedream-5-0-lite-260128"),
            ("豆包视频 · Seedance 2.5", "video", "doubao-seedance-2-5-260628"),
        ]
        legacy_pro_media = session.scalar(select(MediaProvider).where(MediaProvider.name == "豆包配图 · Seedream 5.0 Pro"))
        if legacy_pro_media and not session.scalar(select(MediaProvider).where(MediaProvider.name == "豆包配图 · Seedream 5.0 Lite")):
            legacy_pro_media.name = "豆包配图 · Seedream 5.0 Lite"
            legacy_pro_media.model = "doubao-seedream-5-0-lite-260128"
        for name, media_type, model in media_defaults:
            existing_media = session.scalar(select(MediaProvider).where(MediaProvider.name == name))
            if not existing_media:
                session.add(MediaProvider(name=name, media_type=media_type, model=model, enabled=1 if volcengine and volcengine.api_key_encrypted else 0, api_key_encrypted=volcengine.api_key_encrypted if volcengine else ""))
            elif existing_media.model in {"Doubao-Seedream-5.0-pro", "doubao-seedream-5-0-pro-260628"}:
                existing_media.model = model
        if not session.scalar(select(func.count()).select_from(SearchProvider)):
            session.add(
                SearchProvider(name="Tavily", base_url="https://api.tavily.com", search_depth="advanced", max_results=10, enabled=0)
            )
        else:
            for legacy_search in session.scalars(select(SearchProvider).where(SearchProvider.name == "Tavily")).all():
                legacy_search.enabled = 0
        session.commit()


@app.on_event("startup")
def startup():
    seed()


@app.get("/api/health")
def health():
    return {"status": "online"}


@app.get("/api/dashboard")
def dashboard(session: Session = Depends(db_session)):
    return {"platform_count": 7, "brand_count": session.scalar(select(func.count()).select_from(Brand)), "faq_count": session.scalar(select(func.count()).select_from(Faq)), "knowledge_count": session.scalar(select(func.count()).select_from(Knowledge)), "status": "online"}


@app.get("/api/model-providers")
def list_model_providers(session: Session = Depends(db_session)):
    return [serialize_provider(item) for item in session.scalars(select(ModelProvider).order_by(ModelProvider.id)).all()]


@app.get("/api/search-providers")
def list_search_providers(session: Session = Depends(db_session)):
    return [serialize_search_provider(item) for item in session.scalars(select(SearchProvider).order_by(SearchProvider.id)).all()]


@app.get("/api/media-providers")
def list_media_providers(session: Session = Depends(db_session)):
    return [serialize_media_provider(item) for item in session.scalars(select(MediaProvider).order_by(MediaProvider.media_type, MediaProvider.id)).all()]


@app.post("/api/media-providers")
def create_media_provider(data: MediaProviderInput, session: Session = Depends(db_session)):
    encrypted_key = provider_cipher().encrypt(data.api_key.encode()).decode() if data.api_key else ""
    provider = MediaProvider(**data.model_dump(exclude={"api_key"}), api_key_encrypted=encrypted_key)
    session.add(provider); log(session, "新增媒体模型", provider.name, "新增"); session.commit(); session.refresh(provider)
    return serialize_media_provider(provider)


@app.put("/api/media-providers/{provider_id}")
def update_media_provider(provider_id: int, data: MediaProviderInput, session: Session = Depends(db_session)):
    provider = session.get(MediaProvider, provider_id)
    if not provider:
        raise HTTPException(404, "媒体模型不存在")
    for key, value in data.model_dump(exclude={"api_key"}).items():
        setattr(provider, key, value)
    if data.api_key:
        provider.api_key_encrypted = provider_cipher().encrypt(data.api_key.encode()).decode()
    log(session, "保存媒体模型", provider.name, "修改"); session.commit(); session.refresh(provider)
    return serialize_media_provider(provider)


@app.get("/api/media-providers/{provider_id}/credentials")
def read_media_provider_credentials(provider_id: int, session: Session = Depends(db_session)):
    provider = session.get(MediaProvider, provider_id)
    if not provider:
        raise HTTPException(404, "媒体模型不存在")
    return {"api_key": provider_api_key(provider) if provider.api_key_encrypted else ""}


@app.delete("/api/media-providers/{provider_id}")
def delete_media_provider(provider_id: int, session: Session = Depends(db_session)):
    provider = session.get(MediaProvider, provider_id)
    if not provider:
        raise HTTPException(404, "媒体模型不存在")
    name = provider.name
    session.delete(provider); log(session, "删除媒体模型", name, "删除"); session.commit()
    return {"ok": True}


class MediaGenerateInput(BaseModel):
    provider_id: int
    prompt: str


def extract_media_urls(payload: object) -> list[str]:
    urls: list[str] = []
    def visit(value: object):
        if isinstance(value, dict):
            for key, item in value.items():
                if key in {"url", "image_url", "image_urls", "b64_json", "data"}:
                    if isinstance(item, str) and (item.startswith(("http://", "https://", "data:image/")) or key == "b64_json"):
                        urls.append(item if item.startswith(("http://", "https://", "data:image/")) else f"data:image/png;base64,{item}")
                    elif isinstance(item, list):
                        for child in item:
                            visit(child)
                elif isinstance(item, (dict, list)):
                    visit(item)
        elif isinstance(value, list):
            for item in value:
                visit(item)
    visit(payload)
    return list(dict.fromkeys(urls))


@app.post("/api/media/generate")
def generate_media(data: MediaGenerateInput, session: Session = Depends(db_session)):
    provider = session.get(MediaProvider, data.provider_id)
    if not provider or not provider.enabled:
        raise HTTPException(400, "媒体模型不存在或未启用")
    api_key = provider_api_key(provider) if provider.api_key_encrypted else ""
    if not api_key:
        raise HTTPException(400, "该媒体模型尚未配置 API Key")
    base = provider.base_url.rstrip("/")
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    try:
        if provider.media_type == "image":
            response = httpx.post(f"{base}/images/generations", headers=headers, json={"model": provider.model, "prompt": data.prompt, "n": 3, "response_format": "url"}, timeout=90)
            response.raise_for_status()
            payload = response.json()
            urls = extract_media_urls(payload)
            if not urls:
                raise HTTPException(502, "图片模型未返回可展示的图片 URL，请检查该模型的接口类型、模型名称或火山账户权限")
            return {"media_type": "image", "status": "completed", "urls": urls, "raw": payload}
        response = httpx.post(f"{base}/contents/generations/tasks", headers=headers, json={"model": provider.model, "content": [{"type": "text", "text": data.prompt}]}, timeout=90)
        response.raise_for_status()
        payload = response.json()
        task_id = payload.get("id") or payload.get("task_id")
        return {"media_type": "video", "status": "processing", "task_id": task_id, "raw": payload}
    except httpx.HTTPStatusError as error:
        detail = error.response.text[:500]
        try:
            error_payload = error.response.json().get("error", {})
            if error_payload.get("code") == "SetLimitExceeded":
                detail = "当前账号已达到 Seedream 5.0 Pro 的安全体验模式推理上限。请在火山方舟控制台的模型激活页调整或关闭 Safe Experience Mode，或编辑媒体模型切换到仍有额度的 Seedream Lite。"
        except ValueError:
            pass
        raise HTTPException(error.response.status_code, f"媒体生成接口返回错误：{detail}") from error
    except httpx.HTTPError as error:
        raise HTTPException(502, f"媒体生成请求失败：{error}") from error


@app.get("/api/media/tasks/{task_id}")
def get_media_task(task_id: str, provider_id: int, session: Session = Depends(db_session)):
    provider = session.get(MediaProvider, provider_id)
    if not provider or provider.media_type != "video" or not provider.enabled or not provider.api_key_encrypted:
        raise HTTPException(400, "视频模型不存在、未启用或尚未配置 API Key")
    headers = {"Authorization": f"Bearer {provider_api_key(provider)}"}
    try:
        response = httpx.get(f"{provider.base_url.rstrip('/')}/contents/generations/tasks/{task_id}", headers=headers, timeout=60)
        response.raise_for_status()
        payload = response.json()
        return {"status": payload.get("status", "processing"), "url": payload.get("content", {}).get("video_url") or payload.get("video_url"), "raw": payload}
    except httpx.HTTPStatusError as error:
        raise HTTPException(error.response.status_code, f"视频任务查询错误：{error.response.text[:500]}") from error


@app.get("/api/media/download")
def download_media(url: str):
    parsed = urlparse(url)
    allowed_host = parsed.hostname and (parsed.hostname.endswith(".volces.com") or parsed.hostname.endswith(".volces.com.cn") or parsed.hostname.endswith(".byteimg.com") or parsed.hostname.endswith(".tos-cn-beijing.com"))
    if parsed.scheme not in {"http", "https"} or not parsed.netloc or not allowed_host:
        raise HTTPException(400, "图片地址无效")
    try:
        response = httpx.get(url, timeout=60, follow_redirects=True)
        response.raise_for_status()
        media_type = response.headers.get("content-type", "image/jpeg").split(";", 1)[0]
        return Response(content=response.content, media_type=media_type, headers={"Content-Disposition": "attachment; filename=generated-image"})
    except httpx.HTTPError as error:
        raise HTTPException(502, f"图片下载失败：{error}") from error


@app.post("/api/model-providers")
def create_model_provider(data: ModelProviderInput, session: Session = Depends(db_session)):
    encrypted_key = provider_cipher().encrypt(data.api_key.encode()).decode() if data.api_key else ""
    encrypted_secret = provider_cipher().encrypt(data.api_secret.encode()).decode() if data.api_secret else ""
    encrypted_aux_key = provider_cipher().encrypt(data.aux_api_key.encode()).decode() if data.aux_api_key else ""
    provider = ModelProvider(**data.model_dump(exclude={"api_key", "api_secret", "aux_api_key"}), api_key_encrypted=encrypted_key, api_secret_encrypted=encrypted_secret, aux_api_key_encrypted=encrypted_aux_key)
    session.add(provider); log(session, "新增模型通道", provider.name, "新增"); session.commit(); session.refresh(provider)
    return serialize_provider(provider)


@app.put("/api/model-providers/{provider_id}")
def update_model_provider(provider_id: int, data: ModelProviderInput, session: Session = Depends(db_session)):
    provider = session.get(ModelProvider, provider_id)
    if not provider:
        raise HTTPException(404, "模型通道不存在")
    for key, value in data.model_dump(exclude={"api_key", "api_secret", "aux_api_key"}).items():
        setattr(provider, key, value)
    if data.api_key:
        provider.api_key_encrypted = provider_cipher().encrypt(data.api_key.encode()).decode()
    if data.api_secret:
        provider.api_secret_encrypted = provider_cipher().encrypt(data.api_secret.encode()).decode()
    if data.aux_api_key:
        provider.aux_api_key_encrypted = provider_cipher().encrypt(data.aux_api_key.encode()).decode()
    log(session, "保存模型通道", provider.name, "修改"); session.commit(); session.refresh(provider)
    return serialize_provider(provider)


@app.get("/api/model-providers/{provider_id}/credentials")
def read_model_provider_credentials(provider_id: int, session: Session = Depends(db_session)):
    provider = session.get(ModelProvider, provider_id)
    if not provider:
        raise HTTPException(404, "模型通道不存在")
    return {
        "api_key": provider_api_key(provider) if provider.api_key_encrypted else "",
        "api_secret": provider_api_secret(provider) if provider.api_secret_encrypted else "",
        "aux_api_key": provider_aux_api_key(provider) if provider.aux_api_key_encrypted else "",
    }


@app.delete("/api/model-providers/{provider_id}")
def delete_model_provider(provider_id: int, session: Session = Depends(db_session)):
    provider = session.get(ModelProvider, provider_id)
    if not provider:
        raise HTTPException(404, "模型通道不存在")
    name = provider.name
    session.delete(provider)
    log(session, "删除模型通道", name, "删除")
    session.commit()
    return {"ok": True}


@app.post("/api/search-providers")
def create_search_provider(data: SearchProviderInput, session: Session = Depends(db_session)):
    encrypted_key = provider_cipher().encrypt(data.api_key.encode()).decode() if data.api_key else ""
    provider = SearchProvider(**data.model_dump(exclude={"api_key"}), api_key_encrypted=encrypted_key)
    session.add(provider); log(session, "新增搜索通道", provider.name, "新增"); session.commit(); session.refresh(provider)
    return serialize_search_provider(provider)


@app.put("/api/search-providers/{provider_id}")
def update_search_provider(provider_id: int, data: SearchProviderInput, session: Session = Depends(db_session)):
    provider = session.get(SearchProvider, provider_id)
    if not provider:
        raise HTTPException(404, "搜索通道不存在")
    for key, value in data.model_dump(exclude={"api_key"}).items():
        setattr(provider, key, value)
    if data.api_key:
        provider.api_key_encrypted = provider_cipher().encrypt(data.api_key.encode()).decode()
    log(session, "保存搜索通道", provider.name, "修改"); session.commit(); session.refresh(provider)
    return serialize_search_provider(provider)


@app.get("/api/official-monitor/platforms")
def list_official_monitor_platforms():
    return official_monitor_platforms()


@app.get("/api/official-monitor/tasks")
def list_official_monitor_tasks(session: Session = Depends(db_session)):
    rows = session.scalars(select(OfficialMonitorTask).order_by(OfficialMonitorTask.created_at.desc(), OfficialMonitorTask.id.desc())).all()
    return [serialize_official_task(row) for row in rows]


@app.get("/api/official-monitor/tasks/{task_id}")
def get_official_monitor_task(task_id: int, session: Session = Depends(db_session)):
    task = session.get(OfficialMonitorTask, task_id)
    if not task:
        raise HTTPException(404, "监测任务不存在")
    answer = session_get_answer(task.id, session)
    sources = session_get_sources(task.id, session)
    runs = session.scalars(select(OfficialMonitorRun).where(OfficialMonitorRun.task_id == task.id).order_by(OfficialMonitorRun.created_at.asc(), OfficialMonitorRun.id.asc())).all()
    grouped: dict[str, int] = {}
    for item in sources:
        grouped[item.domain] = grouped.get(item.domain, 0) + 1
    total = len(sources) or 1
    return {
        **serialize_official_task(task),
        "platform_label": OFFICIAL_PLATFORM_META.get(task.platform, {}).get("label", task.platform),
        "answer": serialize(answer) if answer else None,
        "sources": [serialize(item) for item in sources],
        "domains": [{"domain": domain, "mentions": count, "share": f"{round(count * 100 / total)}%"} for domain, count in sorted(grouped.items(), key=lambda pair: (-pair[1], pair[0]))],
        "runs": [serialize(item) for item in runs],
    }


@app.post("/api/official-monitor/tasks")
def create_official_monitor_tasks(data: OfficialMonitorTaskInput, background_tasks: BackgroundTasks, session: Session = Depends(db_session)):
    created = []
    for platform in data.platforms:
        next_id = (session.scalar(select(func.max(OfficialMonitorTask.id))) or 0) + 1 + len(created)
        task = OfficialMonitorTask(
            task_code=f"MON-{next_id:04d}",
            keyword=data.keyword,
            platform=platform,
            note=data.note.strip(),
        )
        session.add(task)
        created.append(task)
    session.commit()
    for task in created:
        session.refresh(task)
        background_tasks.add_task(execute_official_monitor_task, task.id)
    return [serialize_official_task(task) for task in created]


@app.post("/api/official-monitor/tasks/{task_id}/retry")
def retry_official_monitor_task(task_id: int, background_tasks: BackgroundTasks, session: Session = Depends(db_session)):
    task = session.get(OfficialMonitorTask, task_id)
    if not task:
        raise HTTPException(404, "监测任务不存在")
    task.status = "pending"
    task.error_message = ""
    task.started_at = None
    task.finished_at = None
    session.commit()
    background_tasks.add_task(execute_official_monitor_task, task.id)
    return serialize_official_task(task)


@app.get("/api/official-monitor/tasks/{task_id}/screenshot")
def get_official_monitor_screenshot(task_id: int, session: Session = Depends(db_session)):
    answer = session_get_answer(task_id, session)
    if not answer or not answer.screenshot_path:
        raise HTTPException(404, "该任务没有截图")
    path = Path(answer.screenshot_path).resolve()
    artifact_root = OFFICIAL_MONITOR_ARTIFACT_DIR.resolve()
    if artifact_root not in path.parents or not path.exists():
        raise HTTPException(404, "截图文件不存在")
    return FileResponse(path, media_type="image/png", filename=path.name)


@app.post("/api/official-monitor/platforms/{platform}/open-login")
def open_official_monitor_login(platform: str):
    if platform not in OFFICIAL_PLATFORM_META:
        raise HTTPException(404, "平台不存在")
    launcher = Path(__file__).resolve().parent / "official_monitor_login.py"
    profile_dir = OFFICIAL_MONITOR_PROFILE_DIR / platform
    profile_dir.mkdir(parents=True, exist_ok=True)
    subprocess.Popen(
        [
            sys.executable,
            str(launcher),
            "--platform", platform,
            "--profile-dir", str(profile_dir),
        ],
        cwd=str(Path(__file__).resolve().parent.parent.parent),
        start_new_session=True,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
    )
    return {"ok": True, "message": f"已打开 {OFFICIAL_PLATFORM_META[platform]['label']} 登录浏览器。登录完成后请关闭该浏览器窗口，再重试监测任务。"}


def crud_routes(path: str, model, input_model, label: str):
    @app.get(path)
    def list_items(session: Session = Depends(db_session)):
        ordering = Brand.brand_code.asc() if model is Brand else Faq.faq_code.asc() if model is Faq else model.id.desc()
        return [serialize(item) for item in session.scalars(select(model).order_by(ordering)).all()]

    @app.post(path)
    def create_item(data: input_model, session: Session = Depends(db_session)):
        item = model(**data.model_dump())
        if model is Brand:
            next_id = (session.scalar(select(func.max(Brand.id))) or 0) + 1
            item.brand_code = f"BRD-{next_id:03d}"
        if model is Faq:
            next_id = (session.scalar(select(func.max(Faq.id))) or 0) + 1
            item.faq_code = f"FAQ-{next_id:03d}"
        session.add(item); log(session, f"新增{label}", getattr(item, "name", getattr(item, "question", "")), "新增"); session.commit(); session.refresh(item)
        return serialize(item)

    @app.put(f"{path}/{{item_id}}")
    def update_item(item_id: int, data: input_model, session: Session = Depends(db_session)):
        item = session.get(model, item_id)
        if not item: raise HTTPException(404, f"{label}不存在")
        for key, value in data.model_dump().items(): setattr(item, key, value)
        log(session, f"编辑{label}", getattr(item, "name", getattr(item, "question", "")), "修改"); session.commit(); session.refresh(item)
        return serialize(item)

    @app.delete(f"{path}/{{item_id}}")
    def delete_item(item_id: int, session: Session = Depends(db_session)):
        item = session.get(model, item_id)
        if not item: raise HTTPException(404, f"{label}不存在")
        detail = getattr(item, "name", getattr(item, "question", "")); session.delete(item); log(session, f"删除{label}", detail, "删除"); session.commit()
        return {"ok": True}


crud_routes("/api/brands", Brand, BrandInput, "品牌")
crud_routes("/api/faqs", Faq, FaqInput, "FAQ")
crud_routes("/api/knowledge", Knowledge, KnowledgeInput, "知识条目")


@app.put("/api/brands/{item_id}/constraint")
def update_brand_constraint(item_id: int, data: ConstraintInput, session: Session = Depends(db_session)):
    brand = session.get(Brand, item_id)
    if not brand:
        raise HTTPException(404, "品牌不存在")
    brand.constraint = data.constraint
    log(session, "保存品牌约束", f"{brand.brand_code} · {brand.name}", "修改")
    session.commit(); session.refresh(brand)
    return serialize(brand)


@app.get("/api/logs")
def get_logs(session: Session = Depends(db_session)):
    return [serialize(item) for item in session.scalars(select(AuditLog).order_by(AuditLog.created_at.desc()).limit(100)).all()]


@app.delete("/api/logs")
def clear_logs(session: Session = Depends(db_session)):
    session.query(AuditLog).delete(); session.commit(); return {"ok": True}


@app.post("/api/logs/event")
def create_log_event(data: LogEventInput, session: Session = Depends(db_session)):
    log(session, data.action, data.detail, data.category)
    session.commit()
    return {"ok": True}


def response_output_text(payload: dict) -> str:
    if isinstance(payload.get("output_text"), str):
        return payload["output_text"].strip()
    parts: list[str] = []
    for output in payload.get("output", []):
        for content in output.get("content", []) if isinstance(output, dict) else []:
            if isinstance(content, dict) and isinstance(content.get("text"), str):
                parts.append(content["text"])
    return "\n".join(parts).strip()


def search_with_tavily(provider: SearchProvider, keyword: str) -> list[dict]:
    if not provider.api_key_encrypted:
        raise HTTPException(503, f"{provider.name} 未配置 API Key")
    try:
        api_key = provider_cipher().decrypt(provider.api_key_encrypted.encode()).decode()
    except (InvalidToken, ValueError) as error:
        raise HTTPException(503, f"{provider.name} 的 API Key 无法读取，请重新保存") from error
    body = json.dumps({
        "query": keyword,
        "search_depth": provider.search_depth,
        "max_results": provider.max_results,
        "topic": "general",
        "include_answer": False,
        "include_raw_content": False,
    }).encode("utf-8")
    request = Request(
        provider.base_url.rstrip("/") + "/search",
        data=body,
        method="POST",
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
    )
    ssl_context = ssl.create_default_context(cafile=certifi.where())
    try:
        with urlopen(request, timeout=90, context=ssl_context) as response:
            payload = json.loads(response.read().decode("utf-8"))
    except HTTPError as error:
        detail = ""
        try:
            upstream = json.loads(error.read().decode("utf-8"))
            detail = str(upstream.get("detail") or upstream.get("error") or "")[:180]
        except Exception:
            pass
        suffix = f"：{detail}" if detail else ""
        raise HTTPException(502, f"{provider.name} 搜索请求失败（HTTP {error.code}）{suffix}") from error
    except (URLError, TimeoutError) as error:
        raise HTTPException(502, f"{provider.name} 搜索网络连接失败，请检查 Base URL 或网络") from error
    except json.JSONDecodeError as error:
        raise HTTPException(502, f"{provider.name} 搜索返回格式无法解析") from error
    results = []
    for index, row in enumerate(payload.get("results", []), start=1):
        url = str(row.get("url", "")).strip()
        title = str(row.get("title", "")).strip() or url
        if not url:
            continue
        host = url.split("://", 1)[-1].split("/", 1)[0]
        results.append({
            "id": f"S{index:03d}",
            "title": title,
            "url": url,
            "domain": host,
            "snippet": str(row.get("content", "")).strip(),
            "date": str(row.get("published_date") or "未知"),
        })
    if not results:
        raise HTTPException(502, f"{provider.name} 未返回可用搜索结果")
    return results


def provider_api_key(provider: ModelProvider) -> str:
    if not provider.api_key_encrypted:
        raise HTTPException(503, f"{provider.name} 未配置 API Key")
    try:
        return provider_cipher().decrypt(provider.api_key_encrypted.encode()).decode()
    except (InvalidToken, ValueError) as error:
        raise HTTPException(503, f"{provider.name} 的 API Key 无法读取，请重新保存") from error


def provider_api_secret(provider: ModelProvider) -> str:
    if not provider.api_secret_encrypted:
        raise HTTPException(503, f"{provider.name} 未配置 SecretKey")
    try:
        return provider_cipher().decrypt(provider.api_secret_encrypted.encode()).decode()
    except (InvalidToken, ValueError) as error:
        raise HTTPException(503, f"{provider.name} 的 SecretKey 无法读取，请重新保存") from error


def provider_aux_api_key(provider: ModelProvider) -> str:
    if not provider.aux_api_key_encrypted:
        raise HTTPException(503, f"{provider.name} 未配置 TokenHub API Key")
    try:
        return provider_cipher().decrypt(provider.aux_api_key_encrypted.encode()).decode()
    except (InvalidToken, ValueError) as error:
        raise HTTPException(503, f"{provider.name} 的 TokenHub API Key 无法读取，请重新保存") from error


def geo_prompt(keyword: str) -> str:
    return (
        "请联网检索这个关键词，并像模型官方网页版一样给出完整、自然、信息充分的简体中文回答。\n"
        "回答使用 Markdown 正文：用清晰的小标题和段落组织信息；适合比较的数据可使用 Markdown 表格；"
        "需要时使用项目符号或编号列表。不要把整个回答放进 JSON、代码块或单行字符串。\n"
        "所有事实必须来自本次联网搜索。请启用平台原生引用能力，在相关事实后保留 [1]、[ref_1] 或平台原生引用标记，"
        "同时通过 API 的结构化字段返回引用来源。不要在正文末尾重复输出来源 JSON，也不要编造或补写网址。\n"
        f"关键词：{keyword}"
    )


def provider_request(provider: ModelProvider, keyword: str) -> tuple[str, dict]:
    prompt = geo_prompt(keyword)
    base_url = provider.base_url.rstrip("/")
    if provider.provider_type == "qwen":
        return base_url + "/responses", {
            "model": provider.model,
            "input": prompt,
            "tools": [{"type": "web_search"}],
        }
    if provider.provider_type == "hunyuan":
        if "tencentmaas.com" in base_url.lower():
            return base_url + "/chat/completions", {
                "model": provider.model,
                "messages": [{"role": "user", "content": prompt}],
                "stream": False,
                "web_search_options": {"enable": True, "search_source": "standard"},
                "reasoning_effort": "no_think",
            }
        return base_url + "/", {
            "Model": provider.model,
            "Messages": [{"Role": "user", "Content": prompt}],
            "Stream": False,
            "EnableEnhancement": True,
            "ForceSearchEnhancement": True,
            "SearchInfo": True,
            "Citation": True,
        }
    if provider.provider_type == "volcengine":
        return base_url + "/responses", {
            "model": provider.model,
            # Ark Responses expects message content blocks for web-search calls.
            "input": [{
                "role": "user",
                "content": [{"type": "input_text", "text": prompt}],
            }],
            # Web Search can take longer than a normal completion. Ark recommends
            # streaming so the connection is established immediately and the final
            # completed event still carries the complete response and annotations.
            "stream": True,
            "tools": [{"type": "web_search", "limit": 12, "max_keyword": 5}],
            "max_tool_calls": 3,
        }
    if provider.provider_type == "tencent_search":
        return base_url + "/", {"Query": keyword, "Cnt": 10}
    raise HTTPException(400, f"{provider.name} 的厂商类型暂不支持联网信源解析")


def first_text(row: dict, keys: tuple[str, ...]) -> str:
    for key in keys:
        value = row.get(key)
        if isinstance(value, str) and value.strip():
            return value.strip()
    return ""


def structured_sources(payload: dict) -> list[dict]:
    found: list[dict] = []
    found_by_url: dict[str, dict] = {}

    def visit(value):
        if isinstance(value, list):
            for item in value:
                visit(item)
            return
        if not isinstance(value, dict):
            return
        url = first_text(value, ("url", "Url", "URL", "link", "Link"))
        title = first_text(value, ("title", "Title", "name", "Name"))
        if url:
            parsed = urlparse(url)
            if parsed.scheme in {"http", "https"} and parsed.netloc:
                site_name = first_text(value, ("site_name", "siteName", "SiteName", "site", "source", "Text"))
                candidate = {
                    "title": title or url,
                    "url": url,
                    "site_name": site_name or parsed.netloc,
                    "domain": parsed.netloc.lower(),
                    "snippet": first_text(value, ("snippet", "summary", "content", "Content", "text", "Text")),
                    "date": first_text(value, ("published_date", "publish_time", "date", "Date")) or "未知",
                    "index": value.get("index") or value.get("Index") or value.get("citation_index") or len(found) + 1,
                }
                existing = found_by_url.get(url)
                if existing is None:
                    found.append(candidate)
                    found_by_url[url] = candidate
                else:
                    if existing["title"] == url and candidate["title"] != url:
                        existing["title"] = candidate["title"]
                    if existing["site_name"] == parsed.netloc and candidate["site_name"] != parsed.netloc:
                        existing["site_name"] = candidate["site_name"]
                    if not existing["snippet"] and candidate["snippet"]:
                        existing["snippet"] = candidate["snippet"]
                    if existing["date"] == "未知" and candidate["date"] != "未知":
                        existing["date"] = candidate["date"]
                    if value.get("index") or value.get("Index") or value.get("citation_index"):
                        existing["index"] = candidate["index"]
        for child in value.values():
            if isinstance(child, (dict, list)):
                visit(child)

    visit(payload)
    return found


class PageMetadataParser(HTMLParser):
    def __init__(self):
        super().__init__()
        self.in_title = False
        self.title_parts: list[str] = []
        self.og_title = ""
        self.site_name = ""

    def handle_starttag(self, tag: str, attrs):
        values = {str(key).lower(): str(value) for key, value in attrs if value is not None}
        if tag.lower() == "title":
            self.in_title = True
        if tag.lower() == "meta":
            property_name = values.get("property", "").lower()
            name = values.get("name", "").lower()
            if property_name == "og:title":
                self.og_title = values.get("content", "").strip()
            if property_name == "og:site_name" or name == "application-name":
                self.site_name = values.get("content", "").strip()

    def handle_endtag(self, tag: str):
        if tag.lower() == "title":
            self.in_title = False

    def handle_data(self, data: str):
        if self.in_title:
            self.title_parts.append(data)

    @property
    def title(self) -> str:
        return self.og_title or " ".join(" ".join(self.title_parts).split())


def is_public_source_url(url: str) -> bool:
    parsed = urlparse(url)
    if parsed.scheme not in {"http", "https"} or not parsed.hostname:
        return False
    try:
        addresses = {row[4][0] for row in socket.getaddrinfo(parsed.hostname, parsed.port or 443)}
        return bool(addresses) and all(ipaddress.ip_address(address).is_global for address in addresses)
    except (OSError, ValueError):
        return False


def fetch_source_metadata(source: dict) -> dict:
    result = dict(source)
    if not is_public_source_url(source["url"]):
        return result
    try:
        with httpx.Client(timeout=8.0, follow_redirects=True, verify=certifi.where(), headers={"User-Agent": "Mozilla/5.0 GEO-Source-Metadata/1.0"}) as client:
            response = client.get(source["url"])
            response.raise_for_status()
        content_type = response.headers.get("content-type", "").lower()
        if "html" not in content_type:
            return result
        parser = PageMetadataParser()
        parser.feed(response.text[:500_000])
        if parser.title:
            result["title"] = parser.title[:400]
        if parser.site_name:
            result["site_name"] = parser.site_name[:200]
    except (httpx.HTTPError, UnicodeError):
        pass
    return result


def enrich_source_metadata(sources: list[dict], limit: int = 12) -> list[dict]:
    selected = sources[:limit]
    with ThreadPoolExecutor(max_workers=min(6, len(selected) or 1)) as executor:
        return list(executor.map(fetch_source_metadata, selected))


def provider_answer_text(payload: dict) -> str:
    raw_text = response_output_text(payload)
    if raw_text:
        return normalize_answer_text(raw_text)
    choices = payload.get("choices", [])
    if not choices and isinstance(payload.get("output"), dict):
        choices = payload["output"].get("choices", [])
    if not choices and isinstance(payload.get("Response"), dict):
        choices = payload["Response"].get("Choices", [])
    if choices and isinstance(choices[0], dict):
        message = choices[0].get("message") or choices[0].get("Message") or {}
        if isinstance(message, dict):
            content = message.get("content") or message.get("Content") or ""
            if isinstance(content, str):
                return normalize_answer_text(content)
            if isinstance(content, list):
                parts = [first_text(item, ("text", "content")) for item in content if isinstance(item, dict)]
                return normalize_answer_text("\n".join(part for part in parts if part))
    return ""


def call_deepseek_compose(provider: ModelProvider, data: ComposeInput, brand: Brand | None, faq_rows: list[Faq], knowledge_rows: list[Knowledge], constraint: str) -> str:
    api_key = provider_api_key(provider)
    materials = []
    if faq_rows:
        materials.append("FAQ 参考：" + "；".join(f"{row.question}：{row.answer}" for row in faq_rows))
    if knowledge_rows:
        materials.append("知识库参考：" + "；".join(f"{row.name}：{row.content}" for row in knowledge_rows))
    supplemental = f"\n补充要求：{data.supplemental_prompt.strip()}" if data.supplemental_prompt.strip() else ""
    prompt = (
        f"输入主题：{data.prompt}\n\n品牌：{brand.name if brand else '默认品牌'}\n"
        f"品牌约束：{constraint}\n{chr(10).join(materials)}{supplemental}\n\n"
        "请直接生成可发布的中文文章。使用 Markdown 标题、段落、列表和必要的表格；不要输出 JSON、代码块、分析过程或占位说明。"
    )
    body = {
        "model": provider.model,
        "messages": [
            {"role": "system", "content": "你是专业内容编辑，负责生成结构清晰、事实谨慎、可直接发布的中文内容。"},
            {"role": "user", "content": prompt},
        ],
        "stream": False,
    }
    try:
        with httpx.Client(timeout=120.0, verify=certifi.where(), headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}) as client:
            response = client.post(provider.base_url.rstrip("/") + "/chat/completions", json=body)
        response.raise_for_status()
        summary = provider_answer_text(response.json())
    except httpx.HTTPStatusError as error:
        detail = (error.response.text or "")[:180]
        raise HTTPException(502, f"DeepSeek Flash 请求失败（HTTP {error.response.status_code}）{('：' + detail) if detail else ''}") from error
    except httpx.HTTPError as error:
        raise HTTPException(502, "DeepSeek Flash 网络连接失败，请检查 API Key 或网络") from error
    except json.JSONDecodeError as error:
        raise HTTPException(502, "DeepSeek Flash 返回格式无法解析") from error
    if not summary:
        raise HTTPException(502, "DeepSeek Flash 未返回可用文章内容")
    return normalize_answer_text(summary)


def normalize_answer_text(value: str) -> str:
    text_value = value.strip()
    fenced = re.fullmatch(r"```(?:json|markdown|md)?\s*([\s\S]*?)\s*```", text_value, flags=re.IGNORECASE)
    if fenced:
        text_value = fenced.group(1).strip()
    if text_value.startswith(("{", "[")):
        try:
            parsed = json.loads(text_value)
        except json.JSONDecodeError:
            return text_value
        if isinstance(parsed, dict):
            for key in ("answer", "content", "summary", "text", "response"):
                candidate = parsed.get(key)
                if isinstance(candidate, str) and candidate.strip():
                    return candidate.strip()
        if isinstance(parsed, list) and all(isinstance(item, str) for item in parsed):
            return "\n\n".join(item.strip() for item in parsed if item.strip())
    return text_value


def citation_annotation_counts(payload: dict) -> dict[str, int]:
    counts: dict[str, int] = {}

    def visit(value, citation_context: bool = False):
        if isinstance(value, list):
            for item in value:
                visit(item, citation_context)
            return
        if not isinstance(value, dict):
            return
        type_name = str(value.get("type") or value.get("Type") or "").lower()
        is_citation = citation_context or "citation" in type_name or "annotation" in type_name
        url = first_text(value, ("url", "Url", "URL"))
        has_position = any(key in value for key in ("start_index", "end_index", "StartIndex", "EndIndex"))
        if url and (is_citation or has_position):
            counts[url] = counts.get(url, 0) + 1
        for key, child in value.items():
            if isinstance(child, (dict, list)):
                visit(child, is_citation or str(key).lower() in {"annotations", "citations"})

    visit(payload)
    return counts


def citation_index_counts(answer: str) -> dict[int, int]:
    counts: dict[int, int] = {}
    patterns = (
        r"\[(?:ref[_-]?)?(\d+)\]",
        r"【(?:ref[_-]?)?(\d+)】",
    )
    for pattern in patterns:
        for match in re.finditer(pattern, answer, flags=re.IGNORECASE):
            index = int(match.group(1))
            counts[index] = counts.get(index, 0) + 1
    return counts


def tencent_headers(secret_id: str, secret_key: str, body: bytes, host: str, service: str, action: str, version: str) -> dict[str, str]:
    timestamp = int(time.time())
    date = datetime.fromtimestamp(timestamp, timezone.utc).strftime("%Y-%m-%d")
    action_lower = action.lower()
    canonical_headers = f"content-type:application/json; charset=utf-8\nhost:{host}\nx-tc-action:{action_lower}\n"
    signed_headers = "content-type;host;x-tc-action"
    hashed_payload = hashlib.sha256(body).hexdigest()
    canonical_request = f"POST\n/\n\n{canonical_headers}\n{signed_headers}\n{hashed_payload}"
    credential_scope = f"{date}/{service}/tc3_request"
    string_to_sign = f"TC3-HMAC-SHA256\n{timestamp}\n{credential_scope}\n{hashlib.sha256(canonical_request.encode()).hexdigest()}"
    secret_date = hmac.new(("TC3" + secret_key).encode(), date.encode(), hashlib.sha256).digest()
    secret_service = hmac.new(secret_date, service.encode(), hashlib.sha256).digest()
    secret_signing = hmac.new(secret_service, b"tc3_request", hashlib.sha256).digest()
    signature = hmac.new(secret_signing, string_to_sign.encode(), hashlib.sha256).hexdigest()
    authorization = f"TC3-HMAC-SHA256 Credential={secret_id}/{credential_scope}, SignedHeaders={signed_headers}, Signature={signature}"
    return {
        "Authorization": authorization,
        "Content-Type": "application/json; charset=utf-8",
        "Host": host,
        "X-TC-Action": action,
        "X-TC-Version": version,
        "X-TC-Timestamp": str(timestamp),
    }


def hunyuan_headers(secret_id: str, secret_key: str, body: bytes, host: str) -> dict[str, str]:
    return tencent_headers(secret_id, secret_key, body, host, "hunyuan", "ChatCompletions", "2023-09-01")


def volcengine_stream_payload(response: httpx.Response) -> dict:
    completed: dict | None = None
    fallback_output: list[dict] = []
    for raw_line in response.iter_lines():
        line = raw_line.decode("utf-8") if isinstance(raw_line, bytes) else raw_line
        if not line.startswith("data:"):
            continue
        data = line[5:].strip()
        if not data or data == "[DONE]":
            continue
        try:
            event = json.loads(data)
        except json.JSONDecodeError:
            continue
        if event.get("type") == "response.completed" and isinstance(event.get("response"), dict):
            completed = event["response"]
        elif event.get("type") == "response.output_item.done" and isinstance(event.get("item"), dict):
            fallback_output.append(event["item"])
            # The completed assistant message already contains the final text and
            # URL annotations. Do not wait for trailing bookkeeping events.
            if event["item"].get("type") == "message" and event["item"].get("status") in {None, "completed"}:
                return {"output": fallback_output}
    if completed is not None:
        return completed
    if fallback_output:
        return {"output": fallback_output}
    raise HTTPException(502, "火山方舟 / 豆包 流式响应未返回完成事件")


def geo_result_from_sources(provider: ModelProvider, summary: str, payload: dict, selected_pages: list[dict] | None = None) -> dict:
    selected_pages = selected_pages if selected_pages is not None else structured_sources(payload)
    if not selected_pages:
        raise HTTPException(502, f"{provider.name} API 未返回带真实 URL 的结构化信源，请确认该模型已开启联网搜索和来源返回")
    selected_pages = enrich_source_metadata(selected_pages)
    annotation_counts = citation_annotation_counts(payload)
    index_counts = citation_index_counts(summary)
    pages = [
        {
            "title": item["title"], "domain": item["domain"], "site_name": item["site_name"],
            "snippet": item["snippet"],
            "mentions": max(annotation_counts.get(item["url"], 0), index_counts.get(int(item["index"]), 0) if str(item.get("index", "")).isdigit() else 0) or None,
            "date": item["date"], "url": item["url"],
        }
        for item in selected_pages
    ]
    domain_counts: dict[str, int] = {}
    domain_has_unknown: set[str] = set()
    for page in pages:
        if page["mentions"] is None:
            domain_has_unknown.add(page["domain"]); domain_counts.setdefault(page["domain"], 0)
        else:
            domain_counts[page["domain"]] = domain_counts.get(page["domain"], 0) + page["mentions"]
    known_total = sum(domain_counts.values())
    domains = [{"domain": domain, "mentions": count if count or domain not in domain_has_unknown else None, "share": f"{round(count * 100 / known_total)}%" if known_total and count else "未提供"} for domain, count in sorted(domain_counts.items(), key=lambda pair: (-pair[1], pair[0]))]
    return {"model": f"{provider.name} · {provider.model}", "summary": summary, "domains": domains, "pages": pages}


def call_tencent_search_geo(provider: ModelProvider, keyword: str) -> dict:
    secret_id = provider_api_key(provider)
    secret_key = provider_api_secret(provider)
    tokenhub_key = provider_aux_api_key(provider)
    endpoint, request_data = provider_request(provider, keyword)
    request_body = json.dumps(request_data, ensure_ascii=False).encode("utf-8")
    host = urlparse(endpoint).hostname or "wsa.tencentcloudapi.com"
    headers = tencent_headers(secret_id, secret_key, request_body, host, "wsa", "SearchPro", "2025-05-08")
    try:
        with httpx.Client(timeout=90.0, verify=certifi.where(), headers=headers) as client:
            response = client.post(endpoint, content=request_body)
        response.raise_for_status(); search_payload = response.json()
    except httpx.HTTPStatusError as error:
        detail = (error.response.text or "")[:180]
        raise HTTPException(502, f"{provider.name} 搜索请求失败（HTTP {error.response.status_code}）{('：' + detail) if detail else ''}") from error
    except httpx.HTTPError as error:
        raise HTTPException(502, f"{provider.name} 搜索网络连接失败，请检查腾讯云服务开通状态") from error
    response_data = search_payload.get("Response", {})
    if isinstance(response_data, dict) and isinstance(response_data.get("Error"), dict):
        error = response_data["Error"]
        raise HTTPException(502, f"{provider.name} 搜索请求失败：{error.get('Code', '')} {error.get('Message', '')}".strip())
    sources: list[dict] = []
    for index, page_json in enumerate(response_data.get("Pages", []) if isinstance(response_data, dict) else [], start=1):
        try:
            row = json.loads(page_json) if isinstance(page_json, str) else page_json
        except json.JSONDecodeError:
            continue
        if not isinstance(row, dict):
            continue
        url = first_text(row, ("url", "Url"))
        if not url:
            continue
        parsed = urlparse(url)
        sources.append({"title": first_text(row, ("title", "Title")) or url, "url": url, "site_name": first_text(row, ("site", "site_name", "Text")) or parsed.netloc, "domain": parsed.netloc.lower(), "snippet": first_text(row, ("content", "passage", "summary")), "date": first_text(row, ("date", "publish_time")) or "未知", "index": index})
    if not sources:
        raise HTTPException(502, f"{provider.name} 未返回可用搜索结果")
    source_context = "\n\n".join(f"[{item['index']}] {item['title']}\n站点：{item['site_name']}\nURL：{item['url']}\n摘要：{item['snippet'][:900]}" for item in sources)
    summary_body = {"model": provider.model, "messages": [{"role": "system", "content": "你是研究助手。只能基于用户提供的检索资料作答，使用清晰的 Markdown 标题、段落和必要的表格；每个事实后使用对应的 [序号] 引用。不要输出 JSON，不要编造 URL。"}, {"role": "user", "content": f"问题：{keyword}\n\n检索资料：\n{source_context}"}], "stream": False}
    try:
        with httpx.Client(timeout=120.0, verify=certifi.where(), headers={"Authorization": f"Bearer {tokenhub_key}", "Content-Type": "application/json"}) as client:
            summary_response = client.post("https://tokenhub-intl.tencentmaas.com/v1/chat/completions", json=summary_body)
        summary_response.raise_for_status(); summary_payload = summary_response.json()
    except httpx.HTTPStatusError as error:
        raise HTTPException(502, f"{provider.name} 的 TokenHub 摘要请求失败（HTTP {error.response.status_code}）：{(error.response.text or '')[:180]}") from error
    except httpx.HTTPError as error:
        raise HTTPException(502, f"{provider.name} 的 TokenHub 摘要网络连接失败") from error
    summary = provider_answer_text(summary_payload)
    if not summary:
        raise HTTPException(502, f"{provider.name} 的 TokenHub 摘要未返回文本")
    return geo_result_from_sources(provider, summary, {"sources": sources}, sources)


def call_provider_geo(provider: ModelProvider, keyword: str) -> dict:
    if provider.provider_type == "tencent_search":
        return call_tencent_search_geo(provider, keyword)
    api_key = provider_api_key(provider)
    endpoint, request_data = provider_request(provider, keyword)
    request_body = json.dumps(request_data, ensure_ascii=False).encode("utf-8")
    if provider.provider_type == "hunyuan" and "tencentmaas.com" not in provider.base_url.lower():
        api_secret = provider_api_secret(provider)
        host = urlparse(endpoint).hostname or "hunyuan.tencentcloudapi.com"
        headers = hunyuan_headers(api_key, api_secret, request_body, host)
    else:
        headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json", "User-Agent": "curl/8.7.1"}
        if provider.provider_type == "volcengine":
            headers["Accept"] = "text/event-stream"
    try:
        with httpx.Client(timeout=httpx.Timeout(180.0, connect=15.0), verify=certifi.where(), headers=headers) as client:
            if provider.provider_type == "volcengine":
                with client.stream("POST", endpoint, content=request_body) as response:
                    response.raise_for_status()
                    payload = volcengine_stream_payload(response)
            else:
                response = client.post(endpoint, content=request_body)
                response.raise_for_status()
                payload = response.json()
    except httpx.HTTPStatusError as error:
        detail = (error.response.text or "")[:180]
        if "error code: 1010" in detail.lower() or "your request was blocked" in detail.lower():
            raise HTTPException(502, f"{provider.name} 上游网关拒绝访问（Cloudflare 1010）。当前服务器 IP 或地区被 {provider.base_url} 拦截，请更换可用通道、联系供应商放行，或改用其他 GPT 服务。") from error
        raise HTTPException(502, f"{provider.name} API 请求失败（HTTP {error.response.status_code}）{('：' + detail) if detail else ''}") from error
    except httpx.HTTPError as error:
        raise HTTPException(502, f"{provider.name} API 网络连接失败，请检查 Base URL 或网络") from error
    except json.JSONDecodeError as error:
        raise HTTPException(502, f"{provider.name} API 返回格式无法解析") from error
    summary = provider_answer_text(payload)
    if provider.provider_type == "hunyuan" and isinstance(payload.get("Response"), dict) and payload["Response"].get("Error"):
        error = payload["Response"]["Error"]
        raise HTTPException(502, f"{provider.name} API 请求失败：{error.get('Code', '')} {error.get('Message', '')}".strip())
    if not summary:
        raise HTTPException(502, f"{provider.name} API 未返回可用文本")
    selected_pages = structured_sources(payload)
    if not selected_pages:
        raise HTTPException(502, f"{provider.name} API 未返回带真实 URL 的结构化信源，请确认该模型已开启联网搜索和来源返回")
    selected_pages = enrich_source_metadata(selected_pages)
    annotation_counts = citation_annotation_counts(payload)
    index_counts = citation_index_counts(summary)

    pages = [
        {
            "title": item["title"],
            "domain": item["domain"],
            "site_name": item["site_name"],
            "snippet": item["snippet"],
            "mentions": max(
                annotation_counts.get(item["url"], 0),
                index_counts.get(int(item["index"]), 0) if str(item.get("index", "")).isdigit() else 0,
            ) or None,
            "date": item["date"],
            "url": item["url"],
        }
        for item in selected_pages
    ]
    domain_counts: dict[str, int] = {}
    domain_has_unknown: set[str] = set()
    for page in pages:
        if page["mentions"] is None:
            domain_has_unknown.add(page["domain"])
            domain_counts.setdefault(page["domain"], 0)
        else:
            domain_counts[page["domain"]] = domain_counts.get(page["domain"], 0) + page["mentions"]
    known_total = sum(domain_counts.values())
    domains = [
        {
            "domain": domain,
            "mentions": count if count or domain not in domain_has_unknown else None,
            "share": f"{round(count * 100 / known_total)}%" if known_total and count else "未提供",
        }
        for domain, count in sorted(domain_counts.items(), key=lambda pair: (-pair[1], pair[0]))
    ]

    return {"model": f"{provider.name} · {provider.model}", "summary": summary, "domains": domains, "pages": pages}


@app.post("/api/geo/evaluate")
def evaluate_geo(data: GeoInput, session: Session = Depends(db_session)):
    provider_rows = session.scalars(select(ModelProvider).where(ModelProvider.id.in_(data.provider_ids))).all()
    providers_by_id = {provider.id: provider for provider in provider_rows}
    providers = [providers_by_id[provider_id] for provider_id in data.provider_ids if provider_id in providers_by_id]
    if len(providers) != len(data.provider_ids):
        raise HTTPException(404, "所选模型通道不存在")
    disabled = [item.name for item in providers if not item.enabled]
    if disabled:
        raise HTTPException(400, f"模型通道未启用：{', '.join(disabled)}")
    unsupported = [item.name for item in providers if item.provider_type not in {"qwen", "hunyuan", "volcengine", "tencent_search"}]
    if unsupported:
        raise HTTPException(400, f"暂不支持这些模型通道：{', '.join(unsupported)}")
    results: dict[str, dict] = {}
    failures: dict[str, dict[str, str]] = {}
    for provider in providers:
        try:
            results[provider.name] = call_provider_geo(provider, data.keyword)
        except HTTPException as error:
            failures[provider.name] = {"model": provider.model, "error": str(error.detail)}
        except Exception:
            failures[provider.name] = {"model": provider.model, "error": "模型调用失败，请检查模型配置或稍后重试。"}
    if not results:
        details = "；".join(f"{name}：{item['error']}" for name, item in failures.items())
        raise HTTPException(502, f"所有已启用模型均未完成评估。{details}")
    names = list(results)
    attempted_names = [provider.name for provider in providers]
    source_count = sum(len(item["domains"]) for item in results.values())
    page_count = sum(len(item["pages"]) for item in results.values())
    mention_hits = sum(1 for item in results.values() if data.keyword in item["summary"] or item["pages"])
    mention_rate = round(mention_hits * 100 / len(results)) if results else 0
    payload = {"keyword": data.keyword, "platforms": names, "attempted_platforms": attempted_names,
               "failures": failures, "mention_rate": mention_rate, "source_count": source_count,
               "page_count": page_count, "results": results}
    payload["aggregate"] = aggregate_geo_results(payload)
    next_id = (session.scalar(select(func.max(GeoReport.id))) or 0) + 1
    report = GeoReport(
        report_code=f"GEO-{datetime.now():%Y%m%d}-{next_id:04d}-{uuid.uuid4().hex[:4].upper()}",
        keyword=data.keyword,
        provider_names=json.dumps(attempted_names, ensure_ascii=False),
        result_json=json.dumps(payload, ensure_ascii=False),
    )
    session.add(report)
    session.flush()
    failure_note = f" · 未完成:{', '.join(failures)}" if failures else ""
    log(session, "GEO 查询", f"{data.keyword} · 厂商原生联网模型:{', '.join(names)}{failure_note} · {report.report_code}", "查询")
    session.commit()
    session.refresh(report)
    payload["report"] = serialize_geo_report(report)
    return payload


@app.get("/api/geo/reports")
def list_geo_reports(session: Session = Depends(db_session)):
    reports = session.scalars(select(GeoReport).order_by(GeoReport.created_at.desc(), GeoReport.id.desc()).limit(100)).all()
    return [serialize_geo_report(report) for report in reports]


@app.get("/api/geo/reports/{report_id}")
def get_geo_report(report_id: int, session: Session = Depends(db_session)):
    return geo_report_payload(geo_report_or_404(session, report_id))


@app.put("/api/geo/reports/{report_id}")
def update_geo_report(report_id: int, data: GeoReportEditInput, session: Session = Depends(db_session)):
    report = geo_report_or_404(session, report_id)
    try:
        payload = json.loads(report.result_json)
    except json.JSONDecodeError as error:
        raise HTTPException(500, "报告快照数据损坏，无法保存") from error
    results = payload.get("results") if isinstance(payload.get("results"), dict) else {}
    for platform, summary in data.summaries.items():
        if platform in results and isinstance(results[platform], dict):
            results[platform]["summary"] = summary.strip()
    payload["keyword"] = data.keyword
    payload["results"] = results
    payload["aggregate"] = aggregate_geo_results(payload)
    report.keyword = data.keyword
    report.result_json = json.dumps(payload, ensure_ascii=False)
    log(session, "编辑评估报告", f"{report.report_code} · {data.keyword}", "修改")
    session.commit(); session.refresh(report)
    return geo_report_payload(report)


@app.delete("/api/geo/reports/{report_id}")
def delete_geo_report(report_id: int, session: Session = Depends(db_session)):
    report = geo_report_or_404(session, report_id)
    log(session, "删除 GEO 报告", f"{report.report_code} · {report.keyword}", "删除")
    session.delete(report)
    session.commit()
    return {"ok": True}


@app.get("/api/geo/reports/{report_id}/download.docx")
def download_geo_report_docx(report_id: int, session: Session = Depends(db_session)):
    report = geo_report_or_404(session, report_id)
    return Response(
        content=build_geo_report_docx(report),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=report_download_headers(report, "docx"),
    )


@app.get("/api/geo/reports/{report_id}/download.pdf")
def download_geo_report_pdf(report_id: int, session: Session = Depends(db_session)):
    report = geo_report_or_404(session, report_id)
    return Response(
        content=build_geo_report_pdf(report),
        media_type="application/pdf",
        headers=report_download_headers(report, "pdf"),
    )


@app.post("/api/compose")
def compose(data: ComposeInput, session: Session = Depends(db_session)):
    brand = session.get(Brand, data.brand_id) if data.brand_id else session.scalar(select(Brand).order_by(Brand.id))
    faq_rows = session.scalars(select(Faq).order_by(Faq.id).limit(3)).all() if data.include_faq else []
    knowledge_rows = session.scalars(select(Knowledge).order_by(Knowledge.id).limit(2)).all() if data.include_knowledge else []
    constraint = brand.constraint if brand else DEFAULT_CONSTRAINT
    deepseek = session.scalar(select(ModelProvider).where(ModelProvider.provider_type == "deepseek", ModelProvider.enabled == 1).order_by(ModelProvider.id.asc()))
    if not deepseek or not deepseek.api_key_encrypted:
        raise HTTPException(503, "DeepSeek Flash 通道未配置或未启用")
    content = call_deepseek_compose(deepseek, data, brand, faq_rows, knowledge_rows, constraint)
    next_id = (session.scalar(select(func.max(GeneratedArticle.id))) or 0) + 1
    article = GeneratedArticle(
        article_code=f"GEN-{next_id:04d}", input_text=data.prompt,
        supplemental_prompt=data.supplemental_prompt.strip(), content=content,
        brand_id=brand.id if brand else 0, brand_name=brand.name if brand else "默认品牌",
        include_faq=int(data.include_faq), include_knowledge=int(data.include_knowledge),
    )
    session.add(article)
    log(session, "一键成文", f"{article.article_code} · {data.prompt} · 已注入{article.brand_name}约束", "生成")
    session.commit(); session.refresh(article)
    result = serialize_article(article)
    result["constraint"] = constraint
    result["reference_counts"] = {"faq": len(faq_rows), "knowledge": len(knowledge_rows)}
    return result


@app.get("/api/articles")
def list_articles(session: Session = Depends(db_session)):
    rows = session.scalars(select(GeneratedArticle).order_by(GeneratedArticle.created_at.desc())).all()
    return [serialize_article(row) for row in rows]


def update_article_fields(article: GeneratedArticle, data: ArticleDraftInput):
    article.content = data.content
    article.image_prompt = data.image_prompt
    article.selected_images = json.dumps(data.selected_images, ensure_ascii=False)


@app.put("/api/articles/{article_id}")
def save_article(article_id: int, data: ArticleDraftInput, session: Session = Depends(db_session)):
    article = session.get(GeneratedArticle, article_id)
    if not article:
        raise HTTPException(404, "文章不存在")
    update_article_fields(article, data)
    article.content_saved_at = datetime.now()
    log(session, "保存成文文案", article.article_code, "修改")
    session.commit(); session.refresh(article)
    return serialize_article(article)


@app.post("/api/articles/{article_id}/optimize")
def optimize_article(article_id: int, data: ArticleOptimizeInput, session: Session = Depends(db_session)):
    article = session.get(GeneratedArticle, article_id)
    if not article:
        raise HTTPException(404, "文章不存在")
    provider = session.scalar(select(ModelProvider).where(ModelProvider.provider_type == "deepseek", ModelProvider.enabled == 1).order_by(ModelProvider.id.asc()))
    if not provider or not provider.api_key_encrypted:
        raise HTTPException(503, "DeepSeek Flash 通道未配置或未启用")
    brand = session.get(Brand, article.brand_id) if article.brand_id else None
    draft = ComposeInput(prompt=article.content, supplemental_prompt=f"请在保留原文有效信息的基础上，按以下要求优化文案：{data.instruction}", brand_id=article.brand_id or None, include_faq=bool(article.include_faq), include_knowledge=bool(article.include_knowledge))
    faqs = session.scalars(select(Faq).order_by(Faq.id).limit(3)).all() if article.include_faq else []
    knowledge = session.scalars(select(Knowledge).order_by(Knowledge.id).limit(2)).all() if article.include_knowledge else []
    article.content = call_deepseek_compose(provider, draft, brand, faqs, knowledge, brand.constraint if brand else DEFAULT_CONSTRAINT)
    article.content_saved_at = None
    article.confirmed_at = None
    log(session, "优化成文文案", f"{article.article_code} · {data.instruction[:80]}", "生成")
    session.commit(); session.refresh(article)
    return serialize_article(article)


@app.delete("/api/articles/{article_id}")
def delete_article(article_id: int, session: Session = Depends(db_session)):
    article = session.get(GeneratedArticle, article_id)
    if not article:
        raise HTTPException(404, "文章不存在")
    code = article.article_code
    session.query(TranslationTask).filter(TranslationTask.article_id == article_id).delete()
    session.delete(article); log(session, "删除成文历史", code, "删除"); session.commit()
    return {"ok": True}


@app.get("/api/articles/{article_id}/download-package")
def download_article_package(article_id: int, session: Session = Depends(db_session)):
    article = session.get(GeneratedArticle, article_id)
    if not article:
        raise HTTPException(404, "文章不存在")
    try:
        selected = json.loads(article.selected_images or "[]")
    except json.JSONDecodeError:
        selected = []
    output = BytesIO()
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr(f"{article.article_code}/文章.md", article.content)
        for index, url in enumerate(selected, 1):
            parsed = urlparse(url)
            allowed_host = parsed.hostname and (parsed.hostname.endswith(".volces.com") or parsed.hostname.endswith(".volces.com.cn") or parsed.hostname.endswith(".byteimg.com") or parsed.hostname.endswith(".tos-cn-beijing.com"))
            if not allowed_host:
                continue
            try:
                response = httpx.get(url, timeout=60, follow_redirects=True)
                response.raise_for_status()
                extension = "png" if "png" in response.headers.get("content-type", "") else "jpg"
                archive.writestr(f"{article.article_code}/配图-{index}.{extension}", response.content)
            except httpx.HTTPError:
                continue
    output.seek(0)
    return Response(content=output.getvalue(), media_type="application/zip", headers={"Content-Disposition": f"attachment; filename={article.article_code}-图文素材.zip"})


@app.post("/api/articles/{article_id}/confirm")
def confirm_article(article_id: int, data: ArticleDraftInput, session: Session = Depends(db_session)):
    article = session.get(GeneratedArticle, article_id)
    if not article:
        raise HTTPException(404, "文章不存在")
    if not data.selected_images:
        raise HTTPException(400, "确认使用时至少选择 1 张图片")
    if len(data.selected_images) > 3:
        raise HTTPException(400, "最多选择 3 张图片")
    update_article_fields(article, data)
    article.confirmed_at = datetime.now()
    log(session, "确认使用成品文章", f"{article.article_code} · 已保存 {len(data.selected_images)} 张配图", "生成")
    session.commit(); session.refresh(article)
    return serialize_article(article)


def refresh_translation(task: TranslationTask, session: Session):
    if task.status != "processing":
        return
    elapsed = (datetime.now() - task.created_at).total_seconds()
    stages = [8, 25, 48, 72, 90, 100]
    task.progress = stages[min(int(elapsed), len(stages) - 1)]
    if task.progress == 100:
        task.status = "completed"
        task.completed_at = datetime.now()
        article = session.get(GeneratedArticle, task.article_id)
        log(session, "完成图文翻译", f"{article.article_code if article else task.article_id} · {task.language}", "生成")


@app.post("/api/articles/{article_id}/translations")
def create_translation(article_id: int, data: TranslationInput, session: Session = Depends(db_session)):
    article = session.get(GeneratedArticle, article_id)
    if not article:
        raise HTTPException(404, "文章不存在")
    selected = json.loads(article.selected_images or "[]")
    if not selected:
        raise HTTPException(400, "请先确认文章并至少选择 1 张图片")
    active = session.scalar(select(TranslationTask).where(TranslationTask.article_id == article_id, TranslationTask.status == "processing"))
    if active:
        raise HTTPException(409, "该文章已有翻译任务进行中")
    task = TranslationTask(article_id=article_id, language=data.language)
    session.add(task)
    log(session, "图文翻译", f"{article.article_code} · {data.language}", "生成")
    session.commit(); session.refresh(task)
    return serialize(task)


@app.get("/api/translation-tasks/{task_id}")
def get_translation(task_id: int, session: Session = Depends(db_session)):
    task = session.get(TranslationTask, task_id)
    if not task:
        raise HTTPException(404, "翻译任务不存在")
    refresh_translation(task, session)
    session.commit(); session.refresh(task)
    return serialize(task)


@app.get("/api/translation-tasks")
def list_translations(session: Session = Depends(db_session)):
    tasks = session.scalars(select(TranslationTask).order_by(TranslationTask.created_at.desc())).all()
    latest_by_article: dict[int, TranslationTask] = {}
    for task in tasks:
        if task.article_id not in latest_by_article:
            refresh_translation(task, session)
            latest_by_article[task.article_id] = task
    session.commit()
    return [serialize(task) for task in latest_by_article.values()]


@app.post("/api/translation-tasks/{task_id}/retry")
def retry_translation(task_id: int, session: Session = Depends(db_session)):
    task = session.get(TranslationTask, task_id)
    if not task:
        raise HTTPException(404, "翻译任务不存在")
    if task.status != "failed":
        raise HTTPException(400, "只有失败任务可以重试")
    task.status, task.progress, task.error, task.created_at, task.completed_at = "processing", 8, "", datetime.now(), None
    session.commit(); session.refresh(task)
    return serialize(task)
